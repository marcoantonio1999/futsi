from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
DOCS.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)

TODAY = date.today().strftime("%Y-%m-%d")
ACCENT = colors.HexColor("#14532d")
INK = colors.HexColor("#0f172a")
MUTED = colors.HexColor("#52525b")
BORDER = colors.HexColor("#cbd5e1")
FILL = colors.HexColor("#dcfce7")
SOFT = colors.HexColor("#f8fafc")


def pdf_styles():
    styles = getSampleStyleSheet()
    styles["Title"].fontName = "Helvetica-Bold"
    styles["Title"].fontSize = 21
    styles["Title"].leading = 25
    styles["Title"].textColor = INK
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading1"].fontSize = 15
    styles["Heading1"].leading = 18
    styles["Heading1"].spaceBefore = 13
    styles["Heading1"].spaceAfter = 7
    styles["Heading1"].textColor = ACCENT
    styles["Heading2"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontSize = 12
    styles["Heading2"].leading = 15
    styles["Heading2"].spaceBefore = 9
    styles["Heading2"].spaceAfter = 5
    styles["Heading2"].textColor = INK
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 9.2
    styles["BodyText"].leading = 12
    styles["BodyText"].spaceAfter = 5
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10, textColor=MUTED))
    styles.add(
        ParagraphStyle(
            name="Callout",
            parent=styles["BodyText"],
            backColor=colors.HexColor("#f0fdf4"),
            borderColor=colors.HexColor("#bbf7d0"),
            borderWidth=0.6,
            borderPadding=7,
            spaceBefore=6,
            spaceAfter=8,
        )
    )
    return styles


def p(text: str, style):
    return Paragraph(text, style)


def bullets(items: list[str], style):
    return ListFlowable([ListItem(p(item, style)) for item in items], bulletType="bullet", leftIndent=16)


def nums(items: list[str], style):
    return ListFlowable([ListItem(p(item, style)) for item in items], bulletType="1", leftIndent=18)


def table(data: list[list[str]], widths: list[float]):
    styles = pdf_styles()
    t = Table([[p(str(cell), styles["Small"]) for cell in row] for row in data], colWidths=[w * inch for w in widths], repeatRows=1)
    t.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), FILL),
                ("TEXTCOLOR", (0, 0), (-1, 0), INK),
                ("GRID", (0, 0), (-1, -1), 0.35, BORDER),
                ("VALIGN", (0, 0), (-1, -1), "TOP"),
                ("LEFTPADDING", (0, 0), (-1, -1), 5),
                ("RIGHTPADDING", (0, 0), (-1, -1), 5),
                ("TOPPADDING", (0, 0), (-1, -1), 5),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, SOFT]),
            ]
        )
    )
    return t


def build_pdf(path: Path, title: str, blocks: list[dict]):
    styles = pdf_styles()
    story = [
        p(title, styles["Title"]),
        p(f"Futsi Mini ERP | Actualizado {TODAY}", styles["Small"]),
        Spacer(1, 0.12 * inch),
    ]
    for block in blocks:
        kind = block.get("type", "section")
        if kind == "section":
            story.append(p(block["title"], styles["Heading1"]))
            for text in block.get("paragraphs", []):
                story.append(p(text, styles["BodyText"]))
            if block.get("bullets"):
                story.append(bullets(block["bullets"], styles["BodyText"]))
            if block.get("numbers"):
                story.append(nums(block["numbers"], styles["BodyText"]))
            for tdata, widths in block.get("tables", []):
                story.extend([Spacer(1, 0.05 * inch), table(tdata, widths), Spacer(1, 0.08 * inch)])
        elif kind == "callout":
            story.append(p(block["text"], styles["Callout"]))
        elif kind == "pagebreak":
            story.append(PageBreak())
        elif kind == "image":
            image_path = Path(block["path"])
            if image_path.exists():
                story.append(Image(str(image_path), width=block.get("width", 6.8) * inch, height=block.get("height", 3.4) * inch))
                story.append(Spacer(1, 0.1 * inch))

    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.65 * inch, 0.38 * inch, "Futsi Mini ERP - Sprint 2")
        canvas.drawRightString(7.85 * inch, 0.38 * inch, f"Pagina {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def add_docx_table(doc: Document, data: list[list[str]]):
    tbl = doc.add_table(rows=1, cols=len(data[0]))
    tbl.style = "Table Grid"
    hdr = tbl.rows[0].cells
    for idx, cell in enumerate(data[0]):
        hdr[idx].text = str(cell)
    for row in data[1:]:
        cells = tbl.add_row().cells
        for idx, cell in enumerate(row):
            cells[idx].text = str(cell)
    for row in tbl.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)


def build_docx(path: Path, title: str, blocks: list[dict]):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(0.85)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    heading1 = doc.styles["Heading 1"]
    heading1.font.name = "Arial"
    heading1.font.size = Pt(15)
    heading1.font.bold = True
    heading1.font.color.rgb = RGBColor(20, 83, 45)
    heading2 = doc.styles["Heading 2"]
    heading2.font.name = "Arial"
    heading2.font.size = Pt(12)
    heading2.font.bold = True

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title_run = title_p.add_run(title)
    title_run.font.name = "Arial"
    title_run.font.size = Pt(21)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(15, 23, 42)
    meta = doc.add_paragraph(f"Futsi Mini ERP | Actualizado {TODAY}")
    for run in meta.runs:
        run.font.name = "Arial"
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(82, 82, 91)

    for block in blocks:
        kind = block.get("type", "section")
        if kind == "section":
            doc.add_heading(block["title"], level=1)
            for text in block.get("paragraphs", []):
                doc.add_paragraph(text)
            for item in block.get("bullets", []):
                doc.add_paragraph(item, style="List Bullet")
            for item in block.get("numbers", []):
                doc.add_paragraph(item, style="List Number")
            for tdata, _widths in block.get("tables", []):
                add_docx_table(doc, tdata)
                doc.add_paragraph()
        elif kind == "callout":
            para = doc.add_paragraph()
            run = para.add_run(block["text"])
            run.bold = True
            run.font.name = "Arial"
            run.font.size = Pt(9.5)
        elif kind == "pagebreak":
            doc.add_page_break()
    doc.save(path)


def write_both(base_name: str, title: str, blocks: list[dict]):
    pdf_path = DOCS / f"{base_name}.pdf"
    docx_path = DOCS / f"{base_name}.docx"
    build_pdf(pdf_path, title, blocks)
    build_docx(docx_path, title, blocks)
    return pdf_path, docx_path


ROLE_MATRIX = [
    ["Rol", "Responsabilidad", "Acceso Sprint 2"],
    ["Administrador / direccion", "Configura sedes, usuarios, precios, reglas y ve tablero consolidado.", "Dashboard, mapa de sedes, catalogos, reportes, facturacion simulada."],
    ["Contador", "Audita ingresos, egresos, gastos, facturas y exportaciones.", "Reporte contable Excel, graficas, gastos, facturas PDF/XML."],
    ["Coordinador de sede", "Gestiona operacion deportiva y autorizaciones de cancha.", "Asistencia, excepciones, alertas de adeudo, equipo por grupo."],
    ["Coach", "Pasa lista, consulta roster, registra horas y ve alertas medicas/operativas.", "Pase de lista manual y facial, formacion 4-3-3, nomina estimada."],
    ["Cajero / ventanilla", "Procesa cobros sin ver detalles sensibles de direccion.", "Cargos programados, efectivo con aceptacion, transferencia CLABE, terminal/link simulado."],
    ["Padre / tutor", "Consulta adeudos, confirma efectivo, paga por link y descarga facturas.", "Portal familiar, perfil, facturas y notificaciones simuladas."],
]

SPRINT2_CHECKLIST = [
    ["Elemento", "Estado", "Evidencia"],
    ["Excel contable real", "Hecho", "GET /api/reports/accounting.xlsx genera workbook valido con Resumen, Pagos, Cargos, Gastos, Descuentos, Asistencia con adeudo y Facturas."],
    ["Importacion historica Excel", "Hecho base", "Admin/contador cargan Excel, capturan password si aplica, revisan preview editable, firman y generan pagos/gastos historicos con auditoria."],
    ["Facturacion PAC simulada", "Hecho", "Modelo Invoice, endpoint /api/invoices/simulate/, descarga PDF y XML con UUID simulado."],
    ["Facturas en perfiles", "Hecho", "Admin/contador generan; guardian/coach/cajero consultan segun permisos."],
    ["Pase de lista facial", "Hecho local", "DeepFace instalado en .venv; endpoint diagnostico y reconocimiento con recuadro azul/verde/rojo."],
    ["Alumno demo con foto real", "Hecho local", "Marco Antonio Demo en Roma / Equipo Sub-12 A con foto students/photos/retrato_marco.jpeg."],
    ["PWA Android gratis", "Hecho base", "Manifest, icono SVG, service worker y registro en produccion."],
    ["APK Android Capacitor", "Hecho demo", "APK debug generado en front/android/app/build/outputs/apk/debug/app-debug.apk y consume API local via 10.0.2.2."],
    ["Responsive mobile", "Hecho base", "CSS global para header, tablas, formularios, botones tactiles y overflow."],
    ["Tema oscuro web/Android", "Hecho", "Toggle global claro/oscuro con persistencia localStorage, estilos de cards, tablas, inputs, mapa y menu movil."],
    ["Despliegue frontend", "Hecho", "GitHub Pages publica React en /futsi/ con VITE_API_URL a Render."],
    ["Despliegue backend", "Preparado / pausado", "Render + Dockerfile + render.yaml + Supabase Postgres; produccion final se reintentara al cerrar Sprint 2."],
    ["Supabase Postgres", "Preparado", "settings.py soporta POSTGRES_* con sslmode=require para evitar errores por URL malformada."],
]

ANDROID_DELIVERABLES = [
    ["Elemento", "Estado", "Notas"],
    ["PWA instalable", "Hecho", "Gratis, funciona desde Chrome Android, comparte codigo React con web."],
    ["Capacitor Android", "Hecho demo", "Proyecto generado en front/android para empaquetar React como APK."],
    ["APK debug", "Hecho", "front/android/app/build/outputs/apk/debug/app-debug.apk."],
    ["API local en emulador", "Hecho", "Usa VITE_API_URL=http://10.0.2.2:8000/api en modo android."],
    ["Menu movil", "Hecho", "Drawer lateral con cierre de sesion y actualizacion."],
    ["Camara", "Hecho demo", "Captura desde webview; requiere backend local para DeepFace real."],
    ["Tema oscuro", "Hecho", "Toggle global y persistencia local."],
]

SECURITY_SPRINT3_PLAN = [
    ["Area", "Accion Sprint 3", "Criterio de salida"],
    ["SQL injection", "Revisar queries crudas, filtros, reportes y parsers; mantener ORM parametrizado y pruebas con payloads maliciosos.", "Suite automatizada confirma que filtros/login/importacion no ejecutan SQL no esperado."],
    ["Autenticacion", "Validar sesiones/token, expiracion, logout, cambio de password y bloqueo de fuerza bruta.", "Casos negativos cubiertos y logs de intentos fallidos."],
    ["Autorizacion por rol", "Pruebas de acceso horizontal/vertical para admin, contador, cajero, coach, tutor y coordinador.", "Cada endpoint sensible tiene test de permisos."],
    ["Carga de archivos", "Validar extension, MIME, tamano, password, hojas permitidas y limpieza de Excel historico.", "Archivos corruptos/enormes/protegidos fallan de forma controlada."],
    ["Datos personales", "Revisar fotos, datos medicos, responsivas y facturas; minimizar exposicion por rol.", "Cajero/coach/tutor solo ven datos necesarios."],
    ["CORS/CSRF", "Endurecer origenes permitidos, headers, HTTPS y cookies si se cambia estrategia de auth.", "Produccion no acepta origenes comodin."],
    ["Secretos", "Validar Render/GitHub/Supabase sin secretos en repo.", "git-secrets/trufflehog o equivalente sin hallazgos criticos."],
    ["Dependencias", "Ejecutar npm audit/pip-audit y fijar actualizaciones criticas.", "Sin vulnerabilidades altas conocidas o con excepcion documentada."],
]

QA_AUTOMATION_PLAN = [
    ["Herramienta", "Uso propuesto", "Resultado esperado"],
    ["SonarQube / SonarCloud", "Analisis estatico de frontend/backend en GitHub Actions: bugs, smells, cobertura, duplicacion y hotspots.", "Quality Gate visible por PR/main."],
    ["pytest + Django", "Tests unitarios/API para modelos, permisos, pagos, facturas, importacion Excel, reportes y auditoria.", "Suite backend reproducible en CI."],
    ["Playwright Python", "Pruebas e2e usando elementos reales de la web: login por rol, navegacion movil, pagos, asistencia, Excel historico y facturas.", "Flujos criticos pasan en Chrome headless."],
    ["Pruebas agresivas UI", "Scripts que llenan formularios con valores largos, caracteres raros, fechas invalidas, montos negativos, clicks rapidos y reloads.", "La app no se rompe visualmente ni guarda datos invalidos."],
    ["OWASP ZAP baseline", "Escaneo automatizado contra entorno staging local/Render.", "Reporte sin riesgos altos antes de piloto."],
    ["Lighthouse", "Revisar PWA, accesibilidad y performance movil.", "PWA instalable y sin problemas severos de accesibilidad."],
    ["Android smoke", "Instalar APK debug en emulador y validar login, menu, tema oscuro y camara.", "APK usable antes de compartir demo."],
]

SPRINT3_CHECKLIST = [
    ["Pendiente Sprint 3", "Motivo", "Evidencia esperada"],
    ["Hardening de seguridad", "La demo ya toca pagos, menores, fotos, datos medicos y facturas.", "Matriz de amenazas, pruebas de permisos y reporte OWASP/Sonar."],
    ["QA automatizado", "El proyecto ya tiene muchos roles y flujos; probar manualmente no escala.", "CI ejecuta Django tests, Playwright y build web/APK."],
    ["Sonar", "Controlar calidad, duplicacion, bugs y deuda tecnica.", "Quality Gate aprobado en GitHub Actions."],
    ["Importacion historica ampliada", "El primer parser cubre hojas de ingresos/gastos; falta historico completo y staging.", "Excel antiguo cargado en staging, conciliado y firmado."],
    ["Despliegue productivo final", "Sprint 2 quedo local por cambios fuertes.", "Render + Supabase + Pages con smoke test y backup."],
    ["Android estabilizado", "APK debug existe; falta ciclo formal de QA movil.", "APK probado en emulador y celular fisico."],
]

DEPLOYMENT_RESOLUTION = [
    ["Problema", "Causa", "Solucion aplicada"],
    ["GitHub Pages mostraba documentacion o pagina incompleta", "Pages solo sirve estaticos y estaba apuntando a salida/branch incorrecto.", "Workflow de frontend construye Vite y publica front/dist en gh-pages; base path /futsi/."],
    ["Login devolvia Unexpected token '<'", "Frontend llamaba una ruta HTML en lugar del API JSON.", "VITE_API_URL apunta a https://futsi.onrender.com/api y el cliente valida respuesta."],
    ["Backend no podia vivir en Pages", "GitHub Pages no ejecuta procesos Python ni gunicorn.", "Django se hospeda en Render como Web Service; Pages queda solo para React."],
    ["Render fallaba con ValueError port 'H'", "SUPABASE_DATABASE_URL contenia password sin URL encoding, rompiendo parseo de puerto.", "settings.py prioriza POSTGRES_* separados: host, puerto, usuario, password, db y sslmode."],
    ["Produccion quedo inestable durante cambios de Sprint 2", "Se estaban moviendo modelos, dependencias y endpoints nuevos.", "Decision de gobierno: seguir local hasta cerrar Sprint 2 y desplegar en bloque probado."],
]

HISTORICAL_EXCEL_PROFILE = [
    ["Hoja detectada", "Contenido aparente", "Uso para Sprint 3"],
    ["Estimacion Ventas", "Ventas esperadas por sede y mes.", "Comparar venta esperada contra cobros reales/importados para detectar brechas."],
    ["ARCHIVO DE OPERACION", "Detalle operativo con dia, folio, clave, categoria, equipo, responsable, importe y nomina.", "Fuente primaria para transacciones historicas, pagos, conceptos, responsables y posibles fugas."],
    ["ARCHIVO FILIAL", "Informacion amplia de filial/operacion con muchas columnas.", "Mapeo especial; requiere diccionario de columnas antes de importar."],
    ["ESTADO DE RDOS.", "Estado de resultados mensual.", "Base para reconciliar ingresos, egresos y utilidad historica por mes."],
    ["INGRESOS SEDES", "Ingresos por clave/concepto/sede.", "Normalizar historico de ingresos por sede, concepto y periodo."],
    ["GASTOS SEDES", "Egresos por clave/concepto/sede.", "Normalizar gastos historicos y detectar duplicados, inflados o fuera de sede."],
    ["Hoja6", "Resumen global por mes.", "Validacion cruzada contra ingresos/gastos normalizados."],
]

DATA_MODEL_TABLE = [
    ["Entidad", "Uso", "Relaciones clave"],
    ["users / User", "Roles, permisos, sitio primario, coach_group_name, datos de perfil.", "payments.received_by, expenses.captured_by, invoices.issued_by, face_attempts.captured_by."],
    ["sites", "Sedes con coordenadas, cierre operativo y control por ubicacion.", "students, payments, charges, expenses, invoices."],
    ["students", "Alumno de academia con foto, responsable, categoria, grupo, estado, uniforme, responsiva y notas medicas.", "guardian, site, attendance_records, charges, invoices, face_attempts."],
    ["teams / tournaments / rounds", "Torneos adultos por jornada, doble jornada, liguilla, pago semanal o torneo completo.", "charges y payments pueden asociarse a team."],
    ["charges", "Cobros programados por mensualidad, semanalidad, torneo, uniforme, sancion u otro concepto.", "payments, discounts, invoices."],
    ["payments", "Pagos por efectivo, transferencia CLABE, terminal o link simulado.", "charge, student/team, received_by; actualiza estado del cargo."],
    ["expenses", "Gastos operativos pendientes/aprobados/rechazados con evidencia.", "site, captured_by, approved_by, invoices."],
    ["invoices", "Factura simulada ingreso/egreso con UUID, XML y PDF.", "student, guardian, coach, charge, payment, expense."],
    ["attendance_sessions / attendance_records", "Pase de lista manual, cierre y asistencia contra adeudo.", "session, student, created_by."],
    ["face_recognition_attempts", "Bitacora de reconocimiento facial local.", "session, student, captured_by, confidence, engine."],
    ["audit_logs", "Trazabilidad de acciones sensibles.", "actor, entity, before/after."],
]


def srs_blocks():
    return [
        {
            "type": "callout",
            "text": "Objetivo: construir un mini ERP operativo-financiero para academias y torneos, enfocado en control de ingresos, asistencia, gastos, descuentos y trazabilidad por sede, sin reemplazar Contpaqi ni construir un ERP fiscal completo.",
        },
        {
            "title": "1. Alcance funcional",
            "paragraphs": [
                "El sistema debe permitir operar sedes de futbol con academias infantiles y torneos de adultos. La prioridad sigue siendo el 80/20: saber quien entreno o jugo, quien debia pagar, quien pago, quien recibio el pago, quien aplico descuentos y que gastos se cargaron por sede.",
                "Sprint 2 amplia la demo de Sprint 1 con exportacion contable real, facturacion simulada, PWA para Android, pase de lista facial local y preparacion de despliegue con Pages, Render y Supabase.",
            ],
            "bullets": [
                "Academia: alumnos de 4 a 17 anos con representante, fotografia, responsiva, informacion medica, grupo, sede, estado de inscripcion, uniforme, descuentos y pausas.",
                "Torneos: equipos adultos con representante, jugadores, jornadas semanales, doble jornada, torneos de 12 jornadas mas liguilla y esquemas de cobro semanal o torneo completo.",
                "Cobranza: cargos programados, pagos parciales, mora, efectivo con aceptacion del cliente, transferencia por CLABE simulada, terminal/link de tarjeta simulado.",
                "Gastos: captura por sede, aprobacion contable/direccion, evidencia, categoria, proveedor y trazabilidad.",
                "Reportes: dashboard, graficas de ingresos/egresos/utilidad, adeudos, asistencia con adeudo, gastos pendientes, Excel contable y facturas.",
            ],
        },
        {"title": "2. Roles y permisos", "tables": [(ROLE_MATRIX, [1.3, 2.6, 3.1])]},
        {
            "title": "3. Requerimientos funcionales Sprint 1 y 2",
            "tables": [
                (
                    [
                        ["ID", "Requerimiento", "Estado"],
                        ["RF-01", "Autenticacion por rol y navegacion diferenciada.", "Implementado"],
                        ["RF-02", "CRUD de sedes, representantes, alumnos, usuarios y filtros de alumnos.", "Implementado"],
                        ["RF-03", "Pase de lista manual por sede/grupo/fecha/hora con cierre.", "Implementado"],
                        ["RF-04", "Cobros programados por alumno/equipo con pago parcial y estado pendiente/parcial/pagado.", "Implementado"],
                        ["RF-05", "Flujos simulados de efectivo, transferencia CLABE, tarjeta terminal y link de pago.", "Implementado demo"],
                        ["RF-06", "Gastos operativos con aprobacion/rechazo y evidencia.", "Implementado"],
                        ["RF-07", "Dashboard ejecutivo y contable con graficas de ingresos, egresos y utilidad.", "Implementado"],
                        ["RF-08", "Exportacion contable XLSX valida.", "Implementado Sprint 2"],
                        ["RF-09", "PAC simulado para facturas PDF/XML con UUID.", "Implementado Sprint 2"],
                        ["RF-10", "Pase de lista facial local con DeepFace y bitacora de intentos.", "Implementado local Sprint 2"],
                        ["RF-11", "PWA instalable en Android desde Chrome.", "Implementado base Sprint 2"],
                        ["RF-12", "APK Android con Capacitor para pruebas en emulador/dispositivo.", "Implementado demo Sprint 2"],
                        ["RF-13", "Tema oscuro para web y Android con preferencia persistente.", "Implementado Sprint 2"],
                        ["RF-14", "Carga de Excel historico con preview editable, password, firma y auditoria.", "Implementado base Sprint 2; ampliacion Sprint 3"],
                    ],
                    [0.55, 5.4, 1.05],
                )
            ],
        },
        {
            "title": "4. Requerimientos no funcionales",
            "bullets": [
                "Usabilidad: interfaz minimalista, simple para cancha y ventanilla, con botones tactiles y vistas responsive.",
                "Auditoria: cambios sensibles deben registrar usuario, fecha, entidad y detalle antes/despues cuando aplique.",
                "Seguridad: roles limitan informacion financiera y datos personales; cajero no debe ver reportes de direccion.",
                "Privacidad: fotos, datos medicos, identificaciones y responsivas de menores requieren control por rol.",
                "Disponibilidad: frontend estatico en GitHub Pages; backend en proveedor Python; base Supabase Postgres.",
                "Mantenibilidad: backend Django REST, frontend React/Vite/Tailwind, seed demo reproducible y documentos generados por script.",
                "Seguridad Sprint 3: pruebas contra SQL injection, abuso de roles, cargas maliciosas de Excel, secretos expuestos y dependencias vulnerables.",
                "QA Sprint 3: pruebas automatizadas con Django/pytest, Playwright Python, Sonar y escaneo OWASP ZAP en staging.",
            ],
        },
        {"title": "5. Modelo de datos vigente", "tables": [(DATA_MODEL_TABLE, [1.6, 2.5, 3.1])]},
        {
            "title": "6. Criterios de aceptacion Sprint 2",
            "bullets": [
                "El contador descarga un .xlsx valido y abre hojas contables sin recibir HTML ni links internos.",
                "Admin/contador generan factura simulada de un cargo o gasto y descargan PDF/XML con UUID.",
                "Padre/coach consultan sus facturas desde su perfil segun permisos.",
                "Coach abre camara, ve recuadro de reconocimiento y registra asistencia manual o facial.",
                "El sistema indica si DeepFace real esta activo o si se esta usando demo/mock.",
                "La app se puede instalar como PWA en Android para pruebas sin costo de Play Store.",
                "El APK debug generado con Capacitor permite probar en Android Studio/emulador la misma app React.",
                "El tema oscuro funciona en login, portales y dashboards tanto en web como en Android.",
                "El admin/contador puede cargar un Excel historico, revisar placeholders, corregir y firmar antes de afectar pagos/gastos.",
                "La documentacion explica por que Pages no hospeda Django y como queda Render + Supabase.",
            ],
        },
        {"title": "7. Requerimientos Sprint 3: seguridad y QA", "tables": [(SECURITY_SPRINT3_PLAN, [1.6, 3.7, 1.7]), (QA_AUTOMATION_PLAN, [1.6, 3.4, 2.0])]},
    ]


def governance_blocks():
    return [
        {
            "type": "callout",
            "text": "Decision de gobierno: cerrar Sprint 2 completo en local antes de reintentar despliegue productivo, porque se agregaron modelos, endpoints, facturacion simulada, DeepFace local y cambios de infraestructura.",
        },
        {"title": "1. Estado del producto", "tables": [(SPRINT2_CHECKLIST, [2.0, 1.2, 3.8])]},
        {"title": "2. Solucion de despliegue documentada", "tables": [(DEPLOYMENT_RESOLUTION, [1.8, 2.4, 2.8])]},
        {
            "title": "3. Roadmap de sprints",
            "tables": [
                (
                    [
                        ["Sprint", "Objetivo", "Entregables", "Criterio de salida"],
                        ["Sprint 1", "Fuente confiable de operacion.", "Login, sedes, alumnos, asistencia, pagos, gastos, dashboard, auditoria basica.", "Demo de cruce asistencia vs cobranza vs gastos."],
                        ["Sprint 2", "Contabilidad, despliegue y movilidad demo.", "Excel contable, PAC simulado, PWA, DeepFace local, Render/Supabase preparado.", "Demo e2e local estable y documentos actualizados."],
                        ["Sprint 3", "Produccion controlada, historico y seguridad.", "Migraciones Supabase, deploy Render, Pages final, importacion historica Excel ampliada, QA automatizado, Sonar, OWASP, storage, backups y monitoreo.", "URL productiva estable, historico conciliado, quality gate aprobado y pruebas de seguridad sin hallazgos criticos."],
                        ["Sprint futuro", "Automatizacion avanzada.", "SPEI real, terminal real, WhatsApp, OCR, camaras en sede, inventario y modulo deportivo avanzado.", "Integraciones con proveedores contratados."],
                    ],
                    [0.85, 1.55, 3.0, 1.6],
                )
            ],
        },
        {
            "title": "4. Gantt operativo propuesto",
            "tables": [
                (
                    [
                        ["Semana", "Trabajo", "Responsable", "Dependencias"],
                        ["2026-05-25 a 2026-05-31", "Sprint 1 + cierre funcional demo.", "Desarrollo", "Datos demo y reglas base."],
                        ["2026-06-01 a 2026-06-07", "Sprint 2 local: contabilidad, facturacion simulada, PWA y DeepFace.", "Desarrollo + usuario", "Fotos demo, pruebas de Excel/PDF/XML."],
                        ["2026-06-08 a 2026-06-12", "Hardening y deploy: Supabase, Render, Pages, variables, smoke tests.", "Desarrollo + admin cloud", "Password DB, secretos GitHub/Render."],
                        ["2026-06-10 a 2026-06-14", "Importacion historica Excel: perfilado, mapeo, staging, validacion y conciliacion.", "Desarrollo + contador", "Archivo historico, diccionario de claves y criterios de limpieza."],
                        ["2026-06-12 a 2026-06-16", "Seguridad y QA: Sonar, pytest/Django, Playwright Python, pruebas agresivas de formularios y OWASP ZAP baseline.", "Desarrollo + QA", "Entorno staging, usuarios por rol y datos anonimizados."],
                        ["2026-06-15 a 2026-06-21", "Piloto con usuarios de una sede, APK Android y datos historicos conciliados.", "Operacion + contador", "Usuarios reales, historico importado, reglas aprobadas y APK validado."],
                    ],
                    [1.35, 2.8, 1.4, 1.45],
                )
            ],
        },
        {
            "title": "5. Riesgos y controles",
            "bullets": [
                "DeepFace no debe tratarse como control definitivo en produccion sin consentimiento, pruebas de sesgo, camaras controladas y politicas de datos biometricos.",
                "Supabase debe recibir migraciones solo despues de validar localmente; no ejecutar seed_demo --reset en datos reales.",
                "El password de Postgres no debe guardarse en GitHub; en Render usar POSTGRES_PASSWORD como secreto.",
                "Facturacion actual es simulada; no sustituye PAC real ni CFDI fiscal.",
                "Efectivo requiere cierre de caja y aceptacion del cliente para reducir disputas, pero no elimina totalmente riesgo operativo.",
                "Antes de produccion se deben ejecutar pruebas de seguridad contra inyeccion SQL, acceso indebido por rol, cargas maliciosas y secretos expuestos.",
                "La importacion historica puede contaminar reportes si no pasa por staging, conciliacion y firma del contador/admin.",
            ],
        },
        {"title": "6. Sprint 3: checklist de seguridad, QA y cierre", "tables": [(SPRINT3_CHECKLIST, [1.8, 2.6, 2.6]), (SECURITY_SPRINT3_PLAN, [1.6, 3.7, 1.7]), (QA_AUTOMATION_PLAN, [1.6, 3.4, 2.0])]},
    ]


def code_blocks():
    return [
        {
            "title": "1. Arquitectura",
            "paragraphs": [
                "El repositorio esta dividido en back/ para Django REST Framework y front/ para React, Vite y Tailwind. La base local puede correr SQLite; produccion queda preparada para Supabase Postgres. El frontend se publica como estatico y consume el API por VITE_API_URL.",
            ],
            "bullets": [
                "Backend: Django, DRF, modelos core, serializers, viewsets, endpoints de reportes, facturas y reconocimiento facial.",
                "Frontend: React en un solo entry principal, estilos Tailwind/CSS, roles, dashboards y flujos de caja/familia/coach/contador.",
                "Documentos: generados desde tools/build_sprint2_updated_documents.py para mantener consistencia.",
            ],
        },
        {
            "title": "2. Endpoints agregados o reforzados en Sprint 2",
            "tables": [
                (
                    [
                        ["Endpoint", "Metodo", "Uso"],
                        ["/api/reports/accounting.xlsx", "GET", "Descarga Excel contable valido para admin/owner/accounting."],
                        ["/api/invoices/", "GET", "Lista facturas segun rol."],
                        ["/api/invoices/simulate/", "POST", "Genera factura simulada de cargo o gasto."],
                        ["/api/invoices/{id}/pdf/", "GET", "Descarga PDF simulado."],
                        ["/api/invoices/{id}/xml/", "GET", "Descarga XML simulado."],
                        ["/api/face-attendance/recognize/", "GET", "Diagnostico de motor DeepFace/mock."],
                        ["/api/face-attendance/recognize/", "POST", "Compara foto capturada contra roster y registra asistencia."],
                        ["/api/historical-imports/", "GET", "Lista cargas historicas para admin/contador."],
                        ["/api/historical-imports/preview/", "POST", "Carga Excel, password opcional y genera preview editable."],
                        ["/api/historical-imports/{id}/commit/", "POST", "Firma importacion y crea pagos/gastos historicos."],
                    ],
                    [2.35, 0.8, 3.85],
                )
            ],
        },
        {
            "title": "3. Archivos clave modificados",
            "tables": [
                (
                    [
                        ["Archivo", "Responsabilidad"],
                        ["back/core/models.py", "Invoice y FaceRecognitionAttempt; relaciones con cargos, pagos, gastos, alumnos y sesiones."],
                        ["back/core/views.py", "Excel contable, PAC simulado, PDF/XML, diagnostico y reconocimiento DeepFace."],
                        ["back/core/serializers.py", "Serializers para facturas e intentos faciales."],
                        ["back/futsi_api/settings.py", "Config Postgres/Supabase por variables separadas y proteccion ante URL malformada."],
                        ["front/src/main.tsx", "UI de facturas, Excel, PWA, camara, recuadro facial, perfiles y roles."],
                        ["front/src/styles.css", "Responsive mobile, espejo de camara, estados de reconocimiento y tema oscuro."],
                        ["front/public/*", "Manifest, icono y service worker PWA."],
                        ["front/android/*", "Proyecto Android Capacitor y APK debug para emulador/dispositivo."],
                        ["render.yaml / Dockerfile", "Preparacion de deploy Render con migraciones y variables."],
                    ],
                    [2.2, 4.8],
                )
            ],
        },
        {
            "title": "4. Ejecucion local recomendada",
            "bullets": [
                "Backend normal: cd back; python -m venv .venv; .\\.venv\\Scripts\\python.exe -m pip install -r requirements.txt; .\\.venv\\Scripts\\python.exe manage.py migrate; .\\.venv\\Scripts\\python.exe manage.py seed_demo --reset; .\\.venv\\Scripts\\python.exe manage.py runserver.",
                "Backend con DeepFace local: crear .venv en raiz e instalar requirements + deepface + tf-keras; correr manage.py con --noreload para evitar doble carga.",
                "Frontend: cd front; npm install; npm run dev.",
                "Validacion: manage.py test, npm run build, abrir http://127.0.0.1:5173 y revisar roles admin/contador/coach/cajero/familia.",
                "Android: cd front; npm run build:android; npx cap sync android; abrir front/android en Android Studio o usar tools/build_android_debug_apk.ps1.",
            ],
        },
        {
            "title": "5. Notas de DeepFace",
            "bullets": [
                "DeepFace queda opcional y local porque TensorFlow/DeepFace pesan demasiado para una demo gratis en Render.",
                "El backend compara la captura normal y espejada para compensar camaras frontales.",
                "Si DeepFace esta instalado pero no hay match confiable, el sistema reporta sin coincidencia; ya no inventa asistencia por fallback.",
                "FACE_MATCH_MAX_DISTANCE permite ajustar umbral; default 0.55 para demo.",
            ],
        },
        {"title": "6. Importacion historica Excel implementada", "bullets": [
            "Modelos HistoricalImport y HistoricalImportRow guardan archivo original, usuario que subio, usuario que firma, password usado, resumen, filas y destino creado.",
            "El endpoint preview carga Excel con openpyxl y msoffcrypto-tool cuando hay cifrado, detecta hojas INGRESOS SEDES/GASTOS SEDES y genera placeholders editables.",
            "El endpoint commit exige firma responsable, respeta filas omitidas, crea Payment para ingresos y Expense para egresos, y registra AuditLog.",
            "La ampliacion Sprint 3 debe sumar staging completo para ARCHIVO DE OPERACION, ESTADO DE RDOS. y demas hojas historicas.",
        ]},
        {"title": "7. Seguridad y QA tecnica para Sprint 3", "tables": [(SECURITY_SPRINT3_PLAN, [1.6, 3.7, 1.7]), (QA_AUTOMATION_PLAN, [1.6, 3.4, 2.0])]},
    ]


def business_blocks():
    return [
        {
            "title": "1. Diagnostico operacional",
            "paragraphs": [
                "La necesidad central no es tener mas pantallas, sino una fuente confiable de verdad para ingresos, asistencia, descuentos, gastos y cierres por sede. El negocio opera con academias infantiles y torneos adultos, con 400 equipos aproximados y 1000 alumnos, por lo que la trazabilidad por rol y sede es prioritaria.",
            ],
        },
        {
            "title": "2. Flujos de negocio cubiertos",
            "bullets": [
                "Academia: prueba de 2 clases ampliable a 4 semanas, mensualidad por mes completo, uniforme, mora despues de 10 dias, descuentos homogeneos como 15% hermanos y referido.",
                "Torneos: pago semanal o por torneo completo, jornada como semana de juegos, doble jornada, 12 jornadas mas liguilla, decision de jugar con adeudo por coordinador/cajero con topes futuros.",
                "Caja: cajero procesa pagos sin ver direccion; efectivo requiere aceptacion del cliente, transferencia queda en proceso hasta webhook simulado, tarjeta usa terminal/link simulado.",
                "Contabilidad: gastos se solicitan y aprueban; direccion paga; contador revisa estado de resultados, ninos/equipos pagados/no pagados, Excel y facturas simuladas.",
                "Coach: pase de lista manual/facial, horas trabajadas, formacion de equipo y alertas medicas/adeudos.",
                "Historico: admin/contador cargan Excel cerrado, revisan preview, corrigen placeholders, firman y dejan trazabilidad de la carga.",
                "Tema oscuro: disponible para web y Android como mejora de usabilidad en jornadas de tarde/noche.",
            ],
        },
        {
            "title": "3. Controles contra fugas",
            "tables": [
                (
                    [
                        ["Riesgo", "Control implementado", "Pendiente productivo"],
                        ["Alumno/equipo atendido sin pago", "Asistencia contra cargo/adeudo y alerta al pasar lista.", "Reglas definitivas de bloqueo por avisos."],
                        ["Efectivo no reportado", "Pago queda pendiente de aceptacion por tutor/cliente.", "Corte de caja y conciliacion fisica."],
                        ["Transferencia no conciliada", "CLABE simulada y webhook simulado.", "Proveedor SPEI real y webhooks."],
                        ["Gastos inflados o duplicados", "Estados pendiente/aprobado/rechazado y Excel contable.", "Deteccion automatica de duplicados."],
                        ["Facturas no trazables", "Invoice simulado con UUID, PDF y XML.", "PAC real/CFDI si el negocio decide integrarlo."],
                    ],
                    [1.65, 3.0, 2.35],
                )
            ],
        },
        {
            "title": "4. Propuesta Android sin costo",
            "paragraphs": [
                "La alternativa recomendada para pruebas es PWA instalable desde Chrome Android. No requiere Play Store, permite icono en pantalla de inicio y mantiene una sola base React. En un Sprint posterior se podria empaquetar con Capacitor para acceder con mas control a camara, almacenamiento y notificaciones.",
            ],
            "bullets": [
                "Archivos implementados: manifest.webmanifest, icon.svg, sw.js y registro condicional en produccion.",
                "Uso esperado: coordinador/coach/cajero abre la URL Pages o local, inicia sesion y elige Instalar app desde Chrome.",
                "Limite: notificaciones push reales, camara avanzada y modo offline robusto quedan para fase posterior.",
            ],
        },
        {
            "title": "5. Recomendaciones para piloto",
            "numbers": [
                "Elegir una sede piloto y cargar datos reales controlados de 20 alumnos, 3 grupos, 3 equipos y 10 gastos.",
                "Validar con contador que el Excel cubre columnas necesarias para estado de resultados.",
                "Definir reglas finales de mora, avisos y bloqueo para academia y torneos.",
                "Probar PWA en dos celulares Android y una tablet en cancha.",
                "Reintentar deploy productivo despues de migraciones limpias en Supabase y prueba de humo en Render.",
            ],
        },
    ]


def sprint2_update_blocks():
    return [
        {
            "title": "Resumen ejecutivo de lo agregado hoy",
            "bullets": [
                "Reporte contable Excel corregido para generar archivo .xlsx valido desde backend.",
                "PAC simulado para facturas de ingresos y egresos con UUID, XML y PDF descargable.",
                "Facturas visibles en perfiles de familia, coach/cajero y vistas contables segun rol.",
                "Pase de lista facial local con DeepFace, recuadro visual, camara espejada y alumno demo con foto real.",
                "PWA instalable en Android y mejoras responsive para celular.",
                "APK Android Capacitor generado para pruebas en emulador/dispositivo.",
                "Tema oscuro para web y Android con persistencia.",
                "Importacion historica Excel con preview editable, password opcional, firma y auditoria.",
                "Despliegue documentado: frontend en GitHub Pages, backend Django en Render y datos en Supabase Postgres.",
            ],
        },
        {"title": "Checklist Sprint 2", "tables": [(SPRINT2_CHECKLIST, [2.0, 1.2, 3.8])]},
        {"title": "Resolucion de despliegue", "tables": [(DEPLOYMENT_RESOLUTION, [1.8, 2.4, 2.8])]},
        {"title": "Android Sprint 2", "tables": [(ANDROID_DELIVERABLES, [1.8, 1.2, 4.0])]},
        {
            "title": "Estado actual",
            "paragraphs": [
                "La app queda lista para continuar Sprint 2 en local. Produccion se mantiene pausada hasta cerrar las pruebas e2e y luego se desplegara con Render + Supabase para backend y GitHub Pages para frontend.",
            ],
        },
        {
            "title": "Nuevo pendiente agregado para Sprint 3",
            "paragraphs": [
                "Se agrega como pendiente explicito ampliar la importacion historica de todos los Excel anteriores, usando como primer ejemplo `ARCHIVOS USADOS v2.xlsx`. La base ya permite cargar Excel, password opcional, preview editable y firma; Sprint 3 debe llevarlo a staging completo para ingresos, gastos, egresos, estados de resultados, ventas esperadas, responsables, sedes, conceptos y filas operativas.",
            ],
            "tables": [(HISTORICAL_EXCEL_PROFILE, [1.65, 2.55, 2.8])],
        },
        {"title": "Nuevo pendiente agregado para Sprint 3: seguridad y QA", "tables": [(SECURITY_SPRINT3_PLAN, [1.6, 3.7, 1.7]), (QA_AUTOMATION_PLAN, [1.6, 3.4, 2.0])]},
    ]


def srs_appendix_blocks():
    return [
        {
            "title": "7. Supuestos confirmados para Sprint 2",
            "bullets": [
                "Volumen de referencia: alrededor de 400 equipos y 1000 alumnos de academia.",
                "El coordinador es gerente operativo de una sede; el cajero/auxiliar administrativo captura pagos.",
                "Una jornada equivale a una semana de juegos; puede existir doble jornada.",
                "Los torneos duran 12 jornadas mas liguilla; liguilla equivale a finales.",
                "No existe inscripcion inicial; se cobra mensualidad/semanalidad, uniforme, sanciones y conceptos adicionales.",
                "Periodo de prueba normal: 2 clases; puede extenderse hasta 4 semanas.",
                "Mora sugerida: penalizacion despues de 10 dias, con porcentaje configurable.",
                "Costos por sede son distintos; porcentajes de descuento tienden a ser homogeneos.",
                "Se aceptan pagos parciales y deben reflejar saldo pendiente.",
            ],
        },
        {
            "title": "8. Reglas de negocio a configurar",
            "tables": [
                (
                    [
                        ["Regla", "Valor demo", "Pendiente para produccion"],
                        ["Descuento hermanos", "15%", "Validar si aplica a todos los hermanos o a partir del segundo."],
                        ["Descuento referido", "Catalogo demo", "Definir porcentaje, vigencia y evidencia."],
                        ["Mora academia", "5% despues de 10 dias", "Aprobar porcentaje real y fecha de corte."],
                        ["Pago parcial", "Permitido", "Definir monto minimo y aplicacion a adeudo mas antiguo."],
                        ["Avisos por adeudo", "Papa, alumno, bloqueo partido, bloqueo entrenamiento", "Definir tiempos y responsables."],
                        ["Equipo con adeudo", "Puede jugar con mora y autorizacion", "Definir tope de deuda y bloqueo automatico."],
                        ["Gastos", "Coordinador solicita, contador revisa, direccion paga", "Definir montos maximos y evidencia obligatoria."],
                    ],
                    [1.45, 1.75, 3.8],
                )
            ],
        },
        {
            "title": "9. Interfaces externas",
            "bullets": [
                "Supabase Postgres: almacenamiento relacional productivo.",
                "GitHub Pages: hosting estatico del frontend.",
                "Render: ejecucion Django/gunicorn y migraciones.",
                "Proveedor SPEI futuro: CLABE por cliente/equipo y webhooks reales.",
                "Proveedor tarjeta futuro: terminal fisica o link de pago con conciliacion automatica.",
                "PAC futuro: CFDI real; Sprint 2 solo simula XML/PDF/UUID.",
                "WhatsApp/OCR futuro: posible apoyo si se mantiene flujo actual de comprobantes.",
            ],
        },
        {
            "title": "10. Importacion historica Excel para Sprint 3",
            "paragraphs": [
                "El Sprint 3 debe contemplar la captura completa de archivos Excel anteriores para que el sistema no arranque sin contexto historico. El objetivo no es solo migrar alumnos, sino conservar ingresos, egresos, gastos, estados de resultados, ventas esperadas, responsables, sedes, conceptos y evidencia suficiente para comparar periodos anteriores contra la operacion nueva.",
                "Archivo ejemplo revisado: C:\\Users\\daniel\\Documents\\marco\\mexprod\\ARCHIVOS USADOS v2.xlsx. Contiene 7 hojas y aproximadamente 4.4 MB. La hoja ARCHIVO DE OPERACION reporta el maximo de filas de Excel, por lo que requiere perfilado cuidadoso para distinguir rango usado real contra formato extendido.",
            ],
            "tables": [(HISTORICAL_EXCEL_PROFILE, [1.65, 2.55, 2.8])],
        },
        {
            "title": "11. Exclusiones explicitas",
            "bullets": [
                "No se implementa CFDI fiscal real ni timbrado oficial.",
                "No se integra Contpaqi directamente.",
                "No se implementa SPEI real ni terminal bancaria real.",
                "No se implementa modo offline completo.",
                "No se implementa reconocimiento facial productivo en nube.",
                "No se implementa app nativa publicada en Play Store.",
                "No se automatiza decision final de bloqueo deportivo sin reglas aprobadas.",
            ],
        },
    ]


def governance_appendix_blocks():
    return [
        {
            "title": "6. RACI operativo",
            "tables": [
                (
                    [
                        ["Actividad", "Responsable", "Aprueba", "Consultado", "Informado"],
                        ["Reglas de mora/descuento", "Administrador", "Direccion", "Contador / coordinadores", "Cajeros / familias"],
                        ["Captura de pago", "Cajero", "Sistema / cliente en efectivo", "Contador", "Familia / representante"],
                        ["Autorizacion para jugar con adeudo", "Coordinador", "Coordinador / direccion segun tope", "Cajero", "Contador"],
                        ["Gasto operativo", "Coordinador", "Contador / direccion", "Proveedor", "Direccion"],
                        ["Cierre de caja", "Cajero", "Contador", "Coordinador", "Direccion"],
                        ["Deploy productivo", "Desarrollo", "Direccion tecnica", "Usuario negocio", "Operacion"],
                    ],
                    [1.65, 1.35, 1.55, 1.45, 1.0],
                )
            ],
        },
        {
            "title": "7. Definition of Done para Sprint 2",
            "bullets": [
                "Migraciones locales aplicadas sin errores.",
                "Seed demo corre y deja usuarios/alumnos/cargos/gastos suficientes para demo.",
                "Excel contable abre en Excel/LibreOffice y contiene todas las hojas esperadas.",
                "Factura simulada se guarda en DB y permite descarga PDF/XML.",
                "Coach puede pasar lista manual y facial en local con feedback visual.",
                "PWA instala en Android desde Chrome.",
                "Documentacion SRS, gobernanza, codigo, negocio, despliegue y DeepFace actualizada.",
                "Produccion no se reintenta hasta que el usuario apruebe cierre local de Sprint 2.",
            ],
        },
        {
            "title": "8. Runbook de despliegue Sprint 3",
            "numbers": [
                "Congelar rama main con Sprint 2 probado localmente.",
                "Respaldar base local/demo y confirmar que no se ejecutara seed destructivo en produccion.",
                "Crear/confirmar password de database en Supabase.",
                "Configurar POSTGRES_* y secretos Django en Render.",
                "Perfilar Excel historico y crear diccionario de columnas, claves, sedes y conceptos.",
                "Importar historico primero a tablas staging para validar totales antes de afectar tablas operativas.",
                "Conciliar ingresos, egresos, gastos y utilidad contra ESTADO DE RDOS., INGRESOS SEDES y GASTOS SEDES.",
                "Ejecutar deploy Render y revisar logs de migracion.",
                "Probar /health/ y /api/auth/login/ desde Render.",
                "Actualizar variables GitHub Actions: VITE_API_URL y VITE_BASE_PATH.",
                "Publicar Pages y probar login real desde la URL publica.",
                "Descargar Excel, generar factura simulada, crear pago demo y validar reportes historicos.",
                "Registrar evidencia de smoke test, conciliacion historica y aprobar piloto.",
            ],
        },
        {
            "title": "9. Monitoreo minimo",
            "bullets": [
                "Render logs para errores 500, migraciones y CORS.",
                "Supabase dashboard para conexiones, tablas, crecimiento y consultas lentas.",
                "GitHub Actions para build fallido de frontend/backend.",
                "Bitacora funcional: facturas emitidas, intentos faciales, pagos en proceso y gastos pendientes.",
            ],
        },
        {
            "title": "10. Gobierno de migracion historica",
            "bullets": [
                "Ningun Excel historico debe importarse directo a tablas finales sin pasar por staging.",
                "Cada carga debe registrar archivo, hash, usuario, fecha, hojas procesadas, filas leidas, filas rechazadas y totales por sede/mes.",
                "El contador debe aprobar la conciliacion entre origen Excel y datos normalizados antes de activar reportes historicos.",
                "Las diferencias deben clasificarse como formula rota, duplicado, falta de sede, concepto desconocido, fecha invalida, responsable no identificado o monto inconsistente.",
                "El sistema debe permitir reimportar de forma idempotente sin duplicar pagos/gastos.",
            ],
        },
    ]


def code_appendix_blocks():
    return [
        {
            "title": "6. Contrato de Excel contable",
            "tables": [
                (
                    [
                        ["Hoja", "Contenido"],
                        ["Resumen", "Ingresos confirmados, egresos aprobados, utilidad, gastos pendientes, cargos abiertos por sede."],
                        ["Pagos", "Pago individual con cliente, concepto, metodo, canal, estado, monto, fechas y receptor."],
                        ["Cargos", "Adeudos programados por alumno/equipo, concepto, monto, vencimiento y estado."],
                        ["Gastos", "Gasto por sede, categoria, proveedor, monto, estado, capturo/aprobo."],
                        ["Descuentos", "Solicitudes y aprobaciones de descuento con motivo y usuario."],
                        ["Asistencia con adeudo", "Registros donde un alumno asistio teniendo saldo abierto."],
                        ["Facturas", "UUID, tipo, receptor, concepto, subtotal, impuesto, total y fecha."],
                    ],
                    [1.6, 5.4],
                )
            ],
        },
        {
            "title": "7. Contrato de facturacion simulada",
            "bullets": [
                "source_type=expense genera factura de egreso ligada a Expense.",
                "source_type=charge genera factura de ingreso ligada a Charge, Student y Guardian cuando aplica.",
                "El UUID se genera con uuid4 y no representa timbrado fiscal real.",
                "El XML se guarda como texto en DB; el PDF se guarda como archivo generado por reportlab.",
                "Los roles guardian y coach solo ven facturas asociadas a ellos; cajero ve facturas de su sede.",
            ],
        },
        {
            "title": "8. Pipeline facial",
            "numbers": [
                "Frontend abre camara y muestra video espejo.",
                "Al presionar Pasar lista se captura un frame en canvas.",
                "Frontend envia session, image base64 y opcional student forzado.",
                "Backend obtiene roster por sede/grupo y filtra alumnos activos/prueba/pausa/lesion.",
                "Backend compara contra fotos de alumnos con DeepFace y tambien prueba imagen espejada.",
                "Si distancia <= FACE_MATCH_MAX_DISTANCE registra asistencia presente y bitacora.",
                "Si no hay match, devuelve Sin coincidencia y no inventa asistencia.",
            ],
        },
        {
            "title": "9. Variables de entorno relevantes",
            "tables": [
                (
                    [
                        ["Variable", "Uso"],
                        ["DJANGO_SECRET_KEY", "Clave secreta backend."],
                        ["DJANGO_DEBUG", "false en produccion."],
                        ["DJANGO_ALLOWED_HOSTS", "Hosts Render/dominio."],
                        ["CORS_ALLOWED_ORIGINS", "Origenes Pages/local permitidos."],
                        ["DB_ENGINE", "postgres para Supabase."],
                        ["POSTGRES_DB/USER/PASSWORD/HOST/PORT/SSLMODE", "Conexion Supabase separada para evitar errores de URL."],
                        ["FACE_MATCH_MAX_DISTANCE", "Umbral de reconocimiento facial local."],
                        ["VITE_API_URL", "URL del backend para React."],
                        ["VITE_BASE_PATH", "Base path Pages /futsi/."],
                    ],
                    [2.2, 4.8],
                )
            ],
        },
        {
            "title": "10. Diseno tecnico de importacion historica Sprint 3",
            "bullets": [
                "Crear comando Django `import_historical_excel` o endpoint admin protegido para subir archivos historicos.",
                "Leer hojas con parser estructurado, no con manipulacion de strings; aceptar multiples layouts por hoja.",
                "Guardar archivo original y metadatos de carga en `historical_import_batches`.",
                "Guardar filas crudas normalizadas en staging antes de crear Payment, Expense, Charge o metricas historicas.",
                "Generar reporte de errores descargable: fila, hoja, columna, valor, razon y sugerencia.",
                "Crear llaves naturales para deduplicar: periodo, sede, folio, clave, concepto, monto, fecha y fuente.",
                "Exponer dashboard de conciliacion: total origen vs total importado vs diferencia por hoja/sede/mes.",
            ],
        },
    ]


def business_appendix_blocks():
    return [
        {
            "title": "6. KPIs sugeridos",
            "tables": [
                (
                    [
                        ["KPI", "Pregunta que responde", "Fuente"],
                        ["Ingreso por sede", "Cuanto se cobro realmente por sede.", "payments confirmados."],
                        ["Utilidad operativa", "Ingreso menos gasto aprobado.", "payments - expenses approved."],
                        ["Adeudo abierto", "Cuanto falta cobrar.", "charges pending/partial."],
                        ["Asistencia con adeudo", "Quien entreno/jugo sin estar pagado.", "attendance_records + charges."],
                        ["Descuentos por usuario", "Quien aplica o solicita ajustes.", "discounts + audit."],
                        ["Gastos pendientes", "Que gasto todavia no debe afectar resultado.", "expenses pending."],
                        ["Duracion de torneo", "Si el torneo se alarga y pierde rentabilidad.", "rounds/jornadas."],
                        ["Ratio coach/alumno", "Si la nomina deportiva esta proporcionada.", "students + coach logs."],
                    ],
                    [1.6, 2.5, 2.9],
                )
            ],
        },
        {
            "title": "7. Flujo de cobranza objetivo",
            "numbers": [
                "Sistema genera o muestra cargo programado por mensualidad/semanalidad/torneo/uniforme.",
                "Cajero selecciona cliente y metodo de pago.",
                "Transferencia: pago queda en proceso y se simula webhook cuando cae en CLABE del cliente.",
                "Tarjeta: se simula terminal o link, y al completar queda registrado automaticamente.",
                "Efectivo: cajero registra monto y el padre/representante confirma recepcion desde su portal.",
                "Sistema recalcula saldo y actualiza cargo pendiente/parcial/pagado.",
                "Contador revisa Excel, facturas y dashboard.",
            ],
        },
        {
            "title": "8. Flujo de asistencia objetivo",
            "numbers": [
                "Coach o coordinador crea sesion por sede, grupo, fecha y hora.",
                "Sistema lista alumnos del grupo con estado, adeudo, notas medicas y representante.",
                "Se marca asistencia manual o por camara local.",
                "Si existe adeudo o alerta medica, se muestra en pantalla.",
                "Al cerrar sesion, queda trazabilidad y no se edita libremente sin autorizacion futura.",
            ],
        },
        {
            "title": "9. Datos minimos para piloto real",
            "bullets": [
                "Catalogo de sedes, coordenadas, precios y costos por sede.",
                "Lista de alumnos con representante, grupo, telefono, correo, foto y estado.",
                "Lista de equipos, representantes, torneo, tipo de cobro y jornada actual.",
                "Reglas de mora, descuentos, avisos y bloqueo.",
                "Usuarios reales por rol: admin, contador, coordinador, cajero, coach y familias.",
                "Ejemplo de reporte contable esperado y comprobantes/gastos reales.",
            ],
        },
        {
            "title": "10. Valor de negocio del historico",
            "paragraphs": [
                "Importar historico permite que el sistema responda preguntas que no se pueden resolver si solo se captura informacion nueva: que sede tuvo mas diferencia contra ventas esperadas, que conceptos de gasto crecieron sin explicacion, que responsable capturo mas ajustes, donde hubo mas efectivo no conciliado y que torneos o periodos fueron menos rentables.",
            ],
            "bullets": [
                "Comparar ventas esperadas contra ingresos realmente reportados.",
                "Detectar gastos duplicados por sede, concepto, proveedor o fecha.",
                "Medir utilidad historica por sede y mes antes de usar el nuevo sistema.",
                "Identificar sedes con desviaciones recurrentes entre ingresos, egresos y asistencia/operacion.",
                "Dar al contador una base consolidada para dejar de reconstruir reportes largos manualmente.",
            ],
            "tables": [(HISTORICAL_EXCEL_PROFILE, [1.65, 2.55, 2.8])],
        },
        {"title": "11. Android y experiencia movil Sprint 2", "tables": [(ANDROID_DELIVERABLES, [1.8, 1.2, 4.0])]},
        {"title": "12. Seguridad y QA como prioridad Sprint 3", "paragraphs": [
            "El siguiente sprint debe tratar seguridad y QA como entregable funcional, no como tarea secundaria. La app ya maneja dinero, datos de menores, fotos, datos medicos, facturas e historico financiero; por eso el piloto necesita controles antes de operar con datos reales.",
        ], "tables": [(SPRINT3_CHECKLIST, [1.8, 2.6, 2.6])]},
        {
            "title": "13. Flujo de importacion historica Sprint 2/Sprint 3",
            "numbers": [
                "Admin o contador selecciona Excel historico cerrado.",
                "Si el archivo tiene password o cifrado, captura la password en la app.",
                "El sistema analiza hojas soportadas y genera preview editable sin afectar datos finales.",
                "El usuario corrige sede, tipo, fecha, concepto, monto u omite filas incorrectas.",
                "El responsable firma la carga.",
                "El sistema guarda archivo original, quien subio, quien firmo, resumen y filas.",
                "Sprint 2 crea pagos/gastos desde hojas estructuradas; Sprint 3 debe ampliar a staging de todo el historico y conciliacion completa.",
            ],
        },
        {
            "title": "14. Flujo de QA y seguridad Sprint 3",
            "numbers": [
                "Ejecutar Sonar en cada push a main y PR.",
                "Correr pruebas backend Django/pytest contra modelos, permisos, reportes, Excel y facturas.",
                "Correr Playwright Python contra la web con usuarios por rol.",
                "Ejecutar casos agresivos: inputs largos, caracteres especiales, SQL-like payloads, montos negativos, fechas invalidas, doble click y recarga.",
                "Ejecutar OWASP ZAP baseline contra staging.",
                "Documentar hallazgos, prioridad, responsable y evidencia de correccion antes de piloto.",
            ],
        },
    ]


def write_markdown_docs():
    (DOCS / "DEPLOYMENT_SUPABASE.md").write_text(
        """# Despliegue con GitHub Pages, Render y Supabase

## Decision de arquitectura

- Frontend: React/Vite publicado en GitHub Pages.
- Backend: Django REST Framework en Render Web Service.
- Base de datos: Supabase Postgres.
- Archivos: para demo local se usan archivos del backend; para produccion real se recomienda Supabase Storage o volumen persistente.

GitHub Pages no ejecuta Django. Pages solo entrega HTML, CSS y JavaScript estatico. Por eso el backend necesita Render, Railway, Fly.io, DigitalOcean App Platform, Azure App Service o VPS.

## Estado actual

El frontend ya fue preparado para Pages con `VITE_BASE_PATH=/futsi/` y `VITE_API_URL=https://futsi.onrender.com/api`.

El backend quedo preparado para Render y Supabase, pero el despliegue final se pausa hasta cerrar Sprint 2 completo en local. La razon es que hoy se agregaron migraciones, facturacion simulada, reporte Excel y DeepFace local.

## Problemas encontrados y solucion

| Problema | Causa | Solucion |
| --- | --- | --- |
| Pages mostraba documentacion o una pagina estatica incorrecta | Pages estaba publicando otra salida/branch y no `front/dist` | Workflow de frontend publica build Vite en `gh-pages` |
| Login devolvia `Unexpected token '<'` | El frontend recibia HTML en vez de JSON del API | `VITE_API_URL` debe apuntar a `/api` del backend Render |
| Django no puede correr en Pages | Pages no ejecuta procesos Python/gunicorn | Django se despliega en Render como Web Service |
| Render fallaba con `Port could not be cast to integer value as 'H'` | `SUPABASE_DATABASE_URL` tenia password sin URL encoding y rompia el parseo | Usar variables `POSTGRES_*` separadas |

## Variables recomendadas en Render

```env
DJANGO_SECRET_KEY=valor-largo-y-secreto
DJANGO_DEBUG=false
DJANGO_ALLOWED_HOSTS=futsi.onrender.com,.onrender.com
CORS_ALLOWED_ORIGINS=https://marcoantonio1999.github.io,https://marcoantonio1999.github.io/futsi
CSRF_TRUSTED_ORIGINS=https://marcoantonio1999.github.io,https://marcoantonio1999.github.io/futsi
DB_ENGINE=postgres
POSTGRES_DB=postgres
POSTGRES_USER=postgres.uqvjilgskrqehkdpkhvq
POSTGRES_PASSWORD=PASSWORD_REAL_DE_DATABASE
POSTGRES_HOST=aws-1-us-west-2.pooler.supabase.com
POSTGRES_PORT=5432
POSTGRES_SSLMODE=require
DJANGO_SECURE_SSL_REDIRECT=true
DJANGO_SECURE_HSTS_SECONDS=31536000
```

No usar `SUPABASE_DATABASE_URL` si el password tiene caracteres especiales y no esta URL encoded. La configuracion actual de Django prioriza `POSTGRES_*` cuando existen.

## Render

Servicio recomendado: Web Service con Docker.

- Root Directory: vacio.
- Dockerfile Path: `./Dockerfile`.
- Branch: `main`.
- Instance: Free para demo, Starter si se requiere evitar spin down.
- Healthcheck: `/health/`.

El contenedor corre migraciones antes de levantar gunicorn. No ejecutar `seed_demo --reset` en produccion con datos reales.

## GitHub Pages

Variables en GitHub Actions:

```env
VITE_API_URL=https://futsi.onrender.com/api
VITE_BASE_PATH=/futsi/
```

Pages debe publicar desde `gh-pages` o desde el artifact generado por el workflow de frontend.

## Checklist antes de reintentar produccion

- Confirmar password real de base de datos Supabase.
- Configurar secretos en Render.
- Ejecutar migraciones contra Supabase.
- Verificar `/health/` del backend.
- Verificar login desde Pages.
- Descargar Excel contable desde produccion.
- Generar factura simulada PDF/XML.
- Probar CORS entre Pages y Render.
- Decidir almacenamiento productivo para fotos/facturas.
- Perfilar e importar Excel historico `ARCHIVOS USADOS v2.xlsx` a tablas staging.
- Conciliar historico importado contra ingresos, gastos y estado de resultados antes del piloto.
- Ejecutar suite de seguridad Sprint 3: SQL injection, permisos por rol, carga maliciosa de archivos, secretos y dependencias.
- Ejecutar QA automatizado: Sonar, Django/pytest, Playwright Python, OWASP ZAP baseline y smoke Android.
""",
        encoding="utf-8",
    )

    (DOCS / "DEMO_PASE_LISTA_FACIAL.md").write_text(
        """# Demo de pase de lista facial

## Estado Sprint 2

La demo local usa DeepFace cuando esta instalado en el entorno `.venv`. Si DeepFace no esta disponible, el sistema lo muestra en pantalla como `Demo/mock`.

## Funcionalidad implementada

- Camara desde el navegador.
- Video espejado para que se vea natural con camara frontal.
- Recuadro de rostro:
  - Azul: listo.
  - Azul con texto: reconociendo.
  - Verde: coincidencia confiable.
  - Rojo: sin coincidencia.
- Endpoint `GET /api/face-attendance/recognize/` para diagnosticar motor.
- Endpoint `POST /api/face-attendance/recognize/` para enviar captura.
- Bitacora en `face_recognition_attempts`.
- Comparacion de imagen normal y espejada.
- Sin fallback falso cuando DeepFace existe pero no reconoce.

## Alumno demo

Se agrego el alumno `Marco Antonio Demo` en `Roma / Equipo Sub-12 A` con foto local:

`back/media/students/photos/retrato_marco.jpeg`

La fuente original fue:

`C:\\Users\\daniel\\Downloads\\retrato_marco.jpeg`

## Comandos locales usados

```powershell
python -m venv .venv
.\\.venv\\Scripts\\python.exe -m pip install --upgrade pip setuptools wheel
.\\.venv\\Scripts\\python.exe -m pip install -r back\\requirements.txt deepface tf-keras
.\\.venv\\Scripts\\python.exe back\\manage.py runserver 127.0.0.1:8000 --noreload
```

## Validaciones

- `deepface True`
- `tf_keras True`
- `from deepface import DeepFace` correcto
- Foto de Marco contra si misma verificada con confianza 100% en prueba directa.

## Nota de produccion

DeepFace/TensorFlow no se dejan como dependencia obligatoria de Render Free por peso, memoria y tiempo de arranque. Para produccion real se recomienda servicio separado de reconocimiento, GPU o procesamiento asincrono, ademas de consentimiento y politicas de datos biometricos.
""",
        encoding="utf-8",
    )

    (DOCS / "PROPUESTA_ANDROID_GRATIS.md").write_text(
        """# Propuesta Android gratis para Sprint 2

## Recomendacion

Usar la aplicacion como PWA instalable desde Chrome Android. Es gratis, no requiere Play Store y permite probar en cancha con celular o tablet.

Adicionalmente, en Sprint 2 ya se genero una opcion APK con Capacitor para probar en Android Studio, emulador o dispositivo fisico sin publicar en Play Store.

## Implementado

- `front/public/manifest.webmanifest`
- `front/public/icon.svg`
- `front/public/sw.js`
- Registro del service worker en `front/src/main.tsx` solo en produccion.
- Ajustes responsive en `front/src/styles.css`.
- Proyecto Capacitor en `front/android`.
- Script `tools/build_android_debug_apk.ps1`.
- APK debug en `front/android/app/build/outputs/apk/debug/app-debug.apk`.
- Tema oscuro claro/oscuro compartido por web y Android.

## Flujo de uso

1. Abrir la URL de GitHub Pages o la URL local desde Chrome Android.
2. Iniciar sesion con rol coach, cajero, coordinador o familia.
3. En el menu de Chrome elegir Instalar app o Agregar a pantalla principal.
4. Usar la app desde el icono instalado.

## Limitaciones

- El modo offline completo no esta implementado.
- Push notifications reales quedan para fase posterior.
- Para camara avanzada, notificaciones nativas o almacenamiento local robusto se debe estabilizar Capacitor en Sprint 3.

## APK Capacitor

Capacitor ya quedo como demo base. Para regenerarlo:

```powershell
cd front
npm run build
npx cap sync android
cd android
.\\gradlew.bat assembleDebug
```

En emulador Android, la app usa `http://10.0.2.2:8000/api` para consumir el backend local. En produccion debe cambiarse a la URL Render.
""",
        encoding="utf-8",
    )

    (DOCS / "SPRINT3_SEGURIDAD_QA.md").write_text(
        """# Sprint 3 - Seguridad, QA automatizado y hardening

## Objetivo

El Sprint 3 debe cerrar los riesgos principales antes de piloto productivo: pagos, datos de menores, fotos, datos medicos, facturas, historico financiero y permisos por rol.

## Seguridad

| Area | Prueba / control | Salida esperada |
| --- | --- | --- |
| SQL injection | Payloads maliciosos en login, filtros, reportes, importacion Excel y busquedas | ORM parametrizado, sin ejecucion SQL no esperada |
| Autorizacion | Acceso cruzado por rol: cajero, coach, tutor, contador, admin | 403/404 correcto para datos fuera de permiso |
| Archivos Excel | Extension, MIME, tamano, password, cifrado, hojas inesperadas, filas corruptas | Falla controlada y reporte de errores |
| Datos personales | Revisar fotos, responsivas, datos medicos y facturas | Exposicion minima por rol |
| Secretos | Revisar repo, workflows, Render y Supabase | Sin passwords ni tokens en git |
| CORS/CSRF/HTTPS | Origenes permitidos, headers seguros, HTTPS | Produccion sin comodines inseguros |
| Dependencias | npm audit, pip-audit, Dependabot | Vulnerabilidades altas resueltas o documentadas |

## QA automatizado

| Herramienta | Uso |
| --- | --- |
| SonarQube / SonarCloud | Quality Gate en GitHub Actions para bugs, smells, duplicacion, hotspots y cobertura |
| pytest + Django | Modelos, permisos, reportes, facturas, pagos, importacion historica y auditoria |
| Playwright Python | Login por rol, navegacion movil, pagos, asistencia, facturas, historico Excel y tema oscuro |
| Pruebas agresivas UI | Inputs largos, caracteres raros, fechas invalidas, montos negativos, doble click, reload y payloads SQL-like |
| OWASP ZAP baseline | Escaneo contra staging local/Render |
| Lighthouse | PWA, accesibilidad y performance movil |
| Android smoke | Instalar APK debug y probar login, menu, tema oscuro y camara |

## Definition of Done

- GitHub Actions ejecuta build web, check backend, tests y Sonar.
- Playwright cubre al menos admin, contador, cajero, coach y tutor.
- Pruebas de permisos evitan acceso horizontal/vertical.
- ZAP baseline sin hallazgos altos.
- Importacion historica usa staging, conciliacion y firma.
- APK probado en emulador y un celular fisico.
- Deploy Render + Supabase + Pages probado con smoke test documentado.
""",
        encoding="utf-8",
    )


def write_erd():
    erd = """erDiagram
    AUTH_USER ||--o{ PAYMENT : receives
    AUTH_USER ||--o{ EXPENSE : captures
    AUTH_USER ||--o{ INVOICE : issues
    AUTH_USER ||--o{ FACE_RECOGNITION_ATTEMPT : captures
    AUTH_USER ||--o{ HISTORICAL_IMPORT : uploads
    AUTH_USER ||--o{ HISTORICAL_IMPORT : commits

    SITE ||--o{ STUDENT : has
    SITE ||--o{ COURT : has
    SITE ||--o{ CHARGE : bills
    SITE ||--o{ PAYMENT : records
    SITE ||--o{ EXPENSE : spends
    SITE ||--o{ INVOICE : emits
    SITE ||--o{ TOURNAMENT : hosts

    GUARDIAN ||--o{ STUDENT : represents
    GUARDIAN ||--o{ INVOICE : receives

    STUDENT ||--o{ CHARGE : owes
    STUDENT ||--o{ PAYMENT : pays
    STUDENT ||--o{ ATTENDANCE_RECORD : attends
    STUDENT ||--o{ DISCOUNT : gets
    STUDENT ||--o{ INVOICE : receives
    STUDENT ||--o{ FACE_RECOGNITION_ATTEMPT : matches

    TOURNAMENT ||--o{ TEAM : includes
    TOURNAMENT ||--o{ ROUND : has
    TEAM ||--o{ PLAYER : has
    TEAM ||--o{ CHARGE : owes
    TEAM ||--o{ PAYMENT : pays
    TEAM ||--o{ DISCOUNT : gets

    CHARGE ||--o{ PAYMENT : paid_by
    CHARGE ||--o{ DISCOUNT : adjusted_by
    CHARGE ||--o{ INVOICE : invoiced_by
    PAYMENT ||--o{ INVOICE : supports
    EXPENSE ||--o{ INVOICE : invoiced_by

    ATTENDANCE_SESSION ||--o{ ATTENDANCE_RECORD : contains
    ATTENDANCE_SESSION ||--o{ FACE_RECOGNITION_ATTEMPT : has
    HISTORICAL_IMPORT ||--o{ HISTORICAL_IMPORT_ROW : contains

    SITE {
        int id PK
        string name
        string code
        decimal latitude
        decimal longitude
    }
    STUDENT {
        int id PK
        int site_id FK
        int guardian_id FK
        string full_name
        date birth_date
        string group_name
        string status
        string photo
    }
    CHARGE {
        int id PK
        int site_id FK
        int student_id FK
        int team_id FK
        string concept
        decimal amount
        date due_date
        string status
    }
    PAYMENT {
        int id PK
        int charge_id FK
        int student_id FK
        int team_id FK
        string method
        string channel
        string status
        decimal amount
    }
    EXPENSE {
        int id PK
        int site_id FK
        string category
        decimal amount
        string status
    }
    INVOICE {
        int id PK
        uuid uuid
        string kind
        string status
        int charge_id FK
        int payment_id FK
        int expense_id FK
        decimal total
        text xml_content
    }
    FACE_RECOGNITION_ATTEMPT {
        int id PK
        int session_id FK
        int student_id FK
        boolean matched
        decimal confidence
        string engine
    }
    HISTORICAL_IMPORT {
        int id PK
        string original_file
        string original_filename
        string status
        int uploaded_by_id FK
        int committed_by_id FK
        string signature_name
        datetime committed_at
    }
    HISTORICAL_IMPORT_ROW {
        int id PK
        int batch_id FK
        string sheet_name
        int row_number
        json raw_data
        string normalized_type
        string status
    }
"""
    (DOCS / "db_schema_erd.mmd").write_text(erd, encoding="utf-8")


def main():
    write_markdown_docs()
    write_erd()
    outputs = []
    outputs.extend(write_both("Futsi_Especificacion_Requerimientos", "Especificacion de Requerimientos - Sprint 2", srs_blocks() + srs_appendix_blocks()))
    outputs.extend(write_both("Futsi_Gobernanza_Roadmap", "Gobernanza, Roadmap y Despliegue - Sprint 2", governance_blocks() + governance_appendix_blocks()))
    outputs.extend(write_both("Futsi_Documentacion_Codigo", "Documentacion Tecnica del Codigo - Sprint 2", code_blocks() + code_appendix_blocks()))
    outputs.extend(write_both("Futsi_Documentacion_Negocio", "Documentacion de Negocio - Sprint 2", business_blocks() + business_appendix_blocks()))
    outputs.extend(write_both("Futsi_Gobernanza_Sprint2_Actualizacion", "Actualizacion Sprint 2", sprint2_update_blocks() + governance_appendix_blocks()[:2]))
    for output in outputs:
        print(output.relative_to(ROOT))


if __name__ == "__main__":
    main()
