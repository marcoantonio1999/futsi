from __future__ import annotations

import textwrap
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image as RLImage,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
DOCS.mkdir(exist_ok=True)
ASSETS.mkdir(exist_ok=True)


TABLES = {
    "core_user": [
        ("PK", "id", "bigint"),
        ("", "username", "varchar"),
        ("", "email", "varchar"),
        ("", "role", "enum"),
        ("FK", "primary_site_id", "sites.id"),
        ("", "phone", "varchar"),
        ("", "avatar_url", "url"),
        ("", "coach_group_name", "varchar"),
        ("", "coach_hourly_rate", "decimal"),
    ],
    "sites": [
        ("PK", "id", "bigint"),
        ("", "name", "varchar"),
        ("", "code", "slug"),
        ("", "address", "text"),
        ("", "latitude", "decimal"),
        ("", "longitude", "decimal"),
        ("", "is_active", "bool"),
        ("", "close_editing_after_hours", "int"),
    ],
    "courts": [("PK", "id", "bigint"), ("FK", "site_id", "sites.id"), ("", "name", "varchar"), ("", "is_active", "bool")],
    "guardians": [
        ("PK", "id", "bigint"),
        ("FK", "user_id", "core_user.id"),
        ("", "full_name", "varchar"),
        ("", "phone", "varchar"),
        ("", "email", "email"),
        ("", "virtual_clabe", "varchar"),
        ("", "tax_name", "varchar"),
        ("", "tax_id", "varchar"),
    ],
    "students": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "guardian_id", "guardians.id"),
        ("", "full_name", "varchar"),
        ("", "birth_date", "date"),
        ("", "category", "varchar"),
        ("", "group_name", "varchar"),
        ("", "status", "enum"),
        ("", "photo_url", "url"),
        ("", "waiver_url", "url"),
        ("", "medical_notes", "text"),
        ("", "uniform_status", "varchar"),
        ("", "pause_start / pause_end", "date"),
    ],
    "tournaments": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("", "name", "varchar"),
        ("", "billing_type", "enum"),
        ("", "starts_on", "date"),
        ("", "expected_weeks", "int"),
    ],
    "teams": [
        ("PK", "id", "bigint"),
        ("FK", "tournament_id", "tournaments.id"),
        ("", "name", "varchar"),
        ("", "representative_name", "varchar"),
        ("", "representative_phone", "varchar"),
    ],
    "players": [("PK", "id", "bigint"), ("FK", "team_id", "teams.id"), ("", "full_name", "varchar"), ("", "photo", "file"), ("", "identity_document", "file")],
    "rounds": [("PK", "id", "bigint"), ("FK", "tournament_id", "tournaments.id"), ("", "number", "int"), ("", "starts_on / ends_on", "date")],
    "attendance_sessions": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "court_id", "courts.id"),
        ("FK", "tournament_id", "tournaments.id"),
        ("FK", "round_id", "rounds.id"),
        ("FK", "team_id", "teams.id"),
        ("FK", "captured_by_id", "core_user.id"),
        ("", "session_type", "enum"),
        ("", "date / starts_at", "date/time"),
        ("", "group_name", "varchar"),
        ("", "closed_at", "datetime"),
    ],
    "attendance_records": [
        ("PK", "id", "bigint"),
        ("FK", "session_id", "attendance_sessions.id"),
        ("FK", "student_id", "students.id"),
        ("FK", "team_id", "teams.id"),
        ("FK", "captured_by_id", "core_user.id"),
        ("", "status", "enum"),
        ("", "had_debt_at_capture", "bool"),
        ("", "override_reason", "text"),
    ],
    "charges": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "student_id", "students.id"),
        ("FK", "team_id", "teams.id"),
        ("FK", "created_by_id", "core_user.id"),
        ("", "concept", "varchar"),
        ("", "amount", "decimal"),
        ("", "due_date", "date"),
        ("", "status", "enum"),
    ],
    "payments": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "charge_id", "charges.id"),
        ("FK", "student_id", "students.id"),
        ("FK", "team_id", "teams.id"),
        ("FK", "received_by_id", "core_user.id"),
        ("", "method / channel", "enum"),
        ("", "status", "enum"),
        ("", "amount", "decimal"),
        ("", "paid_at / confirmed_at", "datetime"),
        ("", "reference / tracking_key", "varchar"),
        ("", "payment_url", "url"),
    ],
    "discounts": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "charge_id", "charges.id"),
        ("FK", "student_id", "students.id"),
        ("FK", "team_id", "teams.id"),
        ("FK", "requested_by_id", "core_user.id"),
        ("FK", "approved_by_id", "core_user.id"),
        ("", "reason", "varchar"),
        ("", "amount", "decimal"),
        ("", "status", "enum"),
    ],
    "expenses": [
        ("PK", "id", "bigint"),
        ("FK", "site_id", "sites.id"),
        ("FK", "captured_by_id", "core_user.id"),
        ("FK", "approved_by_id", "core_user.id"),
        ("", "category", "varchar"),
        ("", "description", "varchar"),
        ("", "amount", "decimal"),
        ("", "expense_date", "date"),
        ("", "status", "enum"),
    ],
    "coach_work_logs": [
        ("PK", "id", "bigint"),
        ("FK", "coach_id", "core_user.id"),
        ("FK", "site_id", "sites.id"),
        ("FK", "created_by_id", "core_user.id"),
        ("", "group_name", "varchar"),
        ("", "work_date", "date"),
        ("", "hours", "decimal"),
        ("", "hourly_rate_snapshot", "decimal"),
    ],
    "daily_closures": [("PK", "id", "bigint"), ("FK", "site_id", "sites.id"), ("FK", "closed_by_id", "core_user.id"), ("", "business_date", "date"), ("", "cash_expected / cash_reported", "decimal")],
    "audit_logs": [("PK", "id", "bigint"), ("FK", "actor_id", "core_user.id"), ("", "action", "varchar"), ("", "table_name", "varchar"), ("", "record_id", "varchar"), ("", "previous_values / new_values", "json")],
}

RELATIONS = [
    ("core_user", "sites", "primary_site_id"),
    ("guardians", "core_user", "user_id"),
    ("courts", "sites", "site_id"),
    ("students", "sites", "site_id"),
    ("students", "guardians", "guardian_id"),
    ("tournaments", "sites", "site_id"),
    ("teams", "tournaments", "tournament_id"),
    ("players", "teams", "team_id"),
    ("rounds", "tournaments", "tournament_id"),
    ("attendance_sessions", "sites", "site_id"),
    ("attendance_sessions", "courts", "court_id"),
    ("attendance_sessions", "core_user", "captured_by_id"),
    ("attendance_records", "attendance_sessions", "session_id"),
    ("attendance_records", "students", "student_id"),
    ("attendance_records", "teams", "team_id"),
    ("charges", "sites", "site_id"),
    ("charges", "students", "student_id"),
    ("charges", "teams", "team_id"),
    ("payments", "charges", "charge_id"),
    ("payments", "students", "student_id"),
    ("payments", "teams", "team_id"),
    ("discounts", "charges", "charge_id"),
    ("expenses", "sites", "site_id"),
    ("coach_work_logs", "core_user", "coach_id"),
    ("coach_work_logs", "sites", "site_id"),
    ("daily_closures", "sites", "site_id"),
    ("audit_logs", "core_user", "actor_id"),
]


def add_doc_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(15, 23, 42)
    run.bold = True
    sp = doc.add_paragraph(subtitle)
    sp.runs[0].font.name = "Arial"
    sp.runs[0].font.size = Pt(11)
    sp.runs[0].font.color.rgb = RGBColor(82, 82, 91)


def style_doc(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    for name, size, color in [("Heading 1", 15, "0F172A"), ("Heading 2", 12, "14532D"), ("Heading 3", 10.5, "27272A")]:
        style = doc.styles[name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.name = "Arial"
                run.font.size = Pt(8.5)
    for values in rows:
        cells = table.add_row().cells
        for i, value in enumerate(values):
            cells[i].text = str(value)
            for paragraph in cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.name = "Arial"
                    run.font.size = Pt(8)
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)


def bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def create_erd_files() -> Path:
    mermaid_lines = ["erDiagram"]
    for table, fields in TABLES.items():
        mermaid_lines.append(f"  {table} {{")
        for marker, name, dtype in fields:
            key = "PK" if marker == "PK" else "FK" if marker == "FK" else ""
            safe_name = name.replace(" / ", "_").replace(" ", "_").replace(".", "_")
            mermaid_lines.append(f"    {dtype.replace('.', '_')} {safe_name} {key}".rstrip())
        mermaid_lines.append("  }")
    for child, parent, label in RELATIONS:
        mermaid_lines.append(f"  {parent} ||--o{{ {child} : {label}")
    (DOCS / "db_schema_erd.mmd").write_text("\n".join(mermaid_lines), encoding="utf-8")

    svg_parts = [
        '<svg xmlns="http://www.w3.org/2000/svg" width="1800" height="1320" viewBox="0 0 1800 1320">',
        '<rect width="1800" height="1320" fill="#f8fafc"/>',
        '<style>text{font-family:Arial,sans-serif}.t{font-weight:700;font-size:16px;fill:#0f172a}.f{font-size:12px;fill:#334155}.k{font-size:11px;fill:#047857;font-weight:700}.line{stroke:#94a3b8;stroke-width:1.4;fill:none}</style>',
    ]
    positions = {}
    cols = 4
    card_w = 390
    card_h = 222
    x_gap = 40
    y_gap = 38
    for index, (table, fields) in enumerate(TABLES.items()):
        col = index % cols
        row = index // cols
        x = 40 + col * (card_w + x_gap)
        y = 40 + row * (card_h + y_gap)
        positions[table] = (x, y)
        svg_parts.append(f'<rect x="{x}" y="{y}" width="{card_w}" height="{card_h}" rx="8" fill="#ffffff" stroke="#cbd5e1"/>')
        svg_parts.append(f'<rect x="{x}" y="{y}" width="{card_w}" height="34" rx="8" fill="#dcfce7" stroke="#cbd5e1"/>')
        svg_parts.append(f'<text class="t" x="{x+14}" y="{y+23}">{table}</text>')
        for i, (marker, name, dtype) in enumerate(fields[:10]):
            yy = y + 56 + i * 16
            svg_parts.append(f'<text class="k" x="{x+14}" y="{yy}">{marker}</text>')
            svg_parts.append(f'<text class="f" x="{x+48}" y="{yy}">{name}: {dtype}</text>')
        if len(fields) > 10:
            svg_parts.append(f'<text class="f" x="{x+48}" y="{y+56+10*16}">... {len(fields)-10} campos mas</text>')
    for child, parent, label in RELATIONS:
        if child not in positions or parent not in positions:
            continue
        x1, y1 = positions[parent]
        x2, y2 = positions[child]
        start = (x1 + card_w / 2, y1 + card_h)
        end = (x2 + card_w / 2, y2)
        svg_parts.append(f'<path class="line" d="M {start[0]} {start[1]} C {start[0]} {start[1]+20}, {end[0]} {end[1]-20}, {end[0]} {end[1]}"/>')
    svg_parts.append("</svg>")
    svg_path = DOCS / "db_schema_erd.svg"
    svg_path.write_text("\n".join(svg_parts), encoding="utf-8")

    image = Image.new("RGB", (1800, 1320), "#f8fafc")
    draw = ImageDraw.Draw(image)
    try:
        font_title = ImageFont.truetype("arialbd.ttf", 16)
        font_field = ImageFont.truetype("arial.ttf", 12)
        font_key = ImageFont.truetype("arialbd.ttf", 11)
    except OSError:
        font_title = font_field = font_key = ImageFont.load_default()
    for table, fields in TABLES.items():
        x, y = positions[table]
        draw.rounded_rectangle([x, y, x + card_w, y + card_h], radius=8, fill="white", outline="#cbd5e1")
        draw.rounded_rectangle([x, y, x + card_w, y + 34], radius=8, fill="#dcfce7", outline="#cbd5e1")
        draw.text((x + 14, y + 9), table, fill="#0f172a", font=font_title)
        for i, (marker, name, dtype) in enumerate(fields[:10]):
            yy = y + 46 + i * 16
            draw.text((x + 14, yy), marker, fill="#047857", font=font_key)
            draw.text((x + 48, yy), f"{name}: {dtype}", fill="#334155", font=font_field)
        if len(fields) > 10:
            draw.text((x + 48, y + 46 + 10 * 16), f"... {len(fields) - 10} campos mas", fill="#334155", font=font_field)
    png_path = ASSETS / "db_schema_erd.png"
    image.save(png_path)
    return png_path


def create_gantt_image() -> Path:
    tasks = [
        ("Sprint 1", 0, 7, "#047857"),
        ("Dia 1 DB/API base", 0, 1, "#10b981"),
        ("Dia 2 usuarios/sedes", 1, 2, "#10b981"),
        ("Dia 3 asistencia", 2, 3, "#10b981"),
        ("Dia 4 cobranza", 3, 4, "#10b981"),
        ("Dia 5 gastos", 4, 5, "#10b981"),
        ("Dia 6 dashboard/contador", 5, 6, "#10b981"),
        ("Dia 7 QA/documentos", 6, 7, "#10b981"),
        ("Sprint 2", 7, 21, "#0f766e"),
        ("Despliegue staging/prod", 7, 10, "#14b8a6"),
        ("Importacion Excel/Gymforce", 9, 13, "#14b8a6"),
        ("Reglas cobranza/mora", 10, 15, "#14b8a6"),
        ("Conciliacion SPEI real", 13, 18, "#14b8a6"),
        ("Reportes deportivos", 16, 20, "#14b8a6"),
        ("Piloto operacion", 18, 21, "#14b8a6"),
    ]
    w, h = 1200, 620
    image = Image.new("RGB", (w, h), "white")
    draw = ImageDraw.Draw(image)
    try:
        font = ImageFont.truetype("arial.ttf", 15)
        bold = ImageFont.truetype("arialbd.ttf", 16)
    except OSError:
        font = bold = ImageFont.load_default()
    draw.text((32, 24), "Gantt propuesto - Sprint 1 y Sprint 2", fill="#0f172a", font=bold)
    left, top = 260, 80
    day_w = 40
    for day in range(0, 22):
        x = left + day * day_w
        draw.line([x, top - 18, x, h - 40], fill="#e5e7eb")
        if day % 2 == 0:
            draw.text((x - 4, top - 40), str(day), fill="#52525b", font=font)
    for idx, (name, start, end, color) in enumerate(tasks):
        y = top + idx * 32
        draw.text((32, y + 4), name, fill="#27272a", font=font)
        draw.rounded_rectangle([left + start * day_w, y, left + end * day_w, y + 20], radius=5, fill=color)
    path = ASSETS / "gantt_sprint_plan.png"
    image.save(path)
    (DOCS / "gantt_sprint_plan.mmd").write_text(
        "\n".join(
            [
                "gantt",
                "  title Futsi Mini ERP - Sprint 1 y Sprint 2",
                "  dateFormat  YYYY-MM-DD",
                "  section Sprint 1",
                "  DB/API base              :2026-05-26, 1d",
                "  Usuarios/sedes           :2026-05-27, 1d",
                "  Asistencia               :2026-05-28, 1d",
                "  Cobranza                 :2026-05-29, 1d",
                "  Gastos                   :2026-05-30, 1d",
                "  Dashboard/contador       :2026-05-31, 1d",
                "  QA/documentos            :2026-06-01, 1d",
                "  section Sprint 2",
                "  Despliegue staging/prod  :2026-06-02, 3d",
                "  Importacion Excel        :2026-06-04, 4d",
                "  Reglas cobranza/mora     :2026-06-05, 5d",
                "  Conciliacion SPEI real   :2026-06-08, 5d",
                "  Reportes deportivos      :2026-06-11, 4d",
                "  Piloto operacion         :2026-06-13, 3d",
            ]
        ),
        encoding="utf-8",
    )
    return path


def requirements_doc(erd_png: Path) -> Path:
    doc = Document()
    style_doc(doc)
    add_doc_title(doc, "Especificacion de requerimientos - Futsi Mini ERP", "Version Sprint 1, actualizada al 26 de mayo de 2026")
    doc.add_heading("1. Proposito y alcance", level=1)
    doc.add_paragraph("El sistema busca crear una fuente confiable de verdad para sedes, alumnos, representantes, equipos, asistencia, cobranza, descuentos, gastos y reportes operativos. No sustituye Contpaqi ni pretende ser un ERP fiscal completo en Sprint 1.")
    doc.add_heading("2. Diagnostico operativo", level=1)
    bullets(doc, [
        "Excel permite modificaciones posteriores, errores de formula y baja trazabilidad.",
        "Gymforce se usa parcialmente y obliga a exportar o transcribir informacion.",
        "La prioridad es cruzar asistencia, cobranza, descuentos y gastos por sede.",
        "El efectivo, las transferencias no conciliadas y los descuentos sin evidencia son los riesgos principales.",
    ])
    doc.add_heading("3. Roles y permisos", level=1)
    add_table(doc, ["Rol", "Alcance"], [
        ["Administrador", "Configura usuarios, sedes, alumnos, reportes, mapa y parametros."],
        ["Direccion", "Vista ejecutiva consolidada de ingresos, gastos, utilidad y riesgos."],
        ["Contador", "Consulta datos contables, grafica financiera y exporta reporte Excel."],
        ["Coordinador", "Opera sede, asistencia, autorizaciones y seguimiento operativo."],
        ["Cajero", "Procesa pagos de su sede sin ver informacion ejecutiva completa."],
        ["Coach", "Ve su equipo/grupo, pasa lista, registra horas y ve alertas."],
        ["Representante", "Consulta alumnos, adeudos, CLABE, links y confirma efectivo."],
    ], [1.4, 5.0])
    doc.add_heading("4. Requerimientos funcionales Sprint 1", level=1)
    add_table(doc, ["ID", "Requerimiento", "Estado"], [
        ["RF-01", "Login con roles y permisos diferenciados.", "Implementado"],
        ["RF-02", "Catalogo de sedes con coordenadas y mapa real.", "Implementado"],
        ["RF-03", "Registro y edicion de alumnos con datos medicos, foto, responsiva, uniforme y pausas.", "Implementado"],
        ["RF-04", "Representantes con CLABE virtual simulada.", "Implementado"],
        ["RF-05", "Pase de lista por sede/grupo con deuda visible.", "Implementado"],
        ["RF-06", "Cobros programados, pagos parciales y estados de cargo.", "Implementado"],
        ["RF-07", "Simulacion de pagos por transferencia CLABE, efectivo con aceptacion, terminal y link.", "Implementado"],
        ["RF-08", "Gastos por sede con estatus pendiente/aprobado/rechazado.", "Implementado"],
        ["RF-09", "Dashboard de ingresos, egresos, utilidad, adeudos y alertas.", "Implementado"],
        ["RF-10", "Portal del contador con grafica X/Y y exportacion Excel.", "Implementado"],
        ["RF-11", "Portal coach con formacion 4-3-3, banca, asistencia y horas.", "Implementado"],
        ["RF-12", "Bitacora/auditoria amplia de cambios de negocio.", "Parcial"],
    ], [0.7, 4.7, 1.0])
    doc.add_heading("5. Requerimientos no funcionales", level=1)
    bullets(doc, [
        "Interfaz minimalista y clara en Tailwind, pensada para ventanilla y cancha.",
        "API REST en Django/DRF, React/Vite para frontend, SQLite local demo y PostgreSQL recomendado para produccion.",
        "Control de acceso por rol y alcance por sede/grupo/representante.",
        "Proteccion especial de datos de menores: fotos, responsivas, datos medicos e identificaciones.",
        "Exportabilidad hacia Excel para contador y direccion.",
    ])
    doc.add_heading("6. Modelo de datos", level=1)
    doc.add_paragraph("El siguiente diagrama resume entidades principales, llaves primarias y llaves foraneas.")
    doc.add_picture(str(erd_png), width=Inches(6.4))
    doc.add_heading("7. Criterios de aceptacion", level=1)
    bullets(doc, [
        "Un usuario contador puede exportar un archivo Excel con pagos, cargos, gastos, descuentos, resumen y asistencia con adeudo.",
        "Un cajero solo puede operar cobros de su sede.",
        "Un representante ve y confirma solamente pagos de sus propios alumnos.",
        "Un coach solo ve su grupo/equipo asignado.",
        "La demo permite demostrar fuga potencial: alumno/equipo asistio o jugo con adeudo.",
    ])
    path = DOCS / "Futsi_Especificacion_Requerimientos.docx"
    doc.save(path)
    return path


def governance_doc(gantt_png: Path) -> Path:
    doc = Document()
    style_doc(doc)
    add_doc_title(doc, "Gobernanza, roadmap y control de ejecucion", "Sprint 1 + Sprint 2 propuesto para Futsi Mini ERP")
    doc.add_heading("1. Principio de gobernanza", level=1)
    doc.add_paragraph("La prioridad es control y trazabilidad antes que automatizacion compleja. Cada cambio relevante debe dejar actor, fecha, entidad, monto/estado previo y estado nuevo.")
    doc.add_heading("2. Checklist de avance", level=1)
    add_table(doc, ["Area", "Estado actual", "Siguiente paso"], [
        ["Roles", "Admin, direccion, contador, coordinador, cajero, coach y representante.", "Afinar permisos finales con responsables reales."],
        ["Alumnos", "Datos completos, foto, responsiva, medico, uniforme, pausas.", "Importacion masiva desde Excel/Gymforce."],
        ["Asistencia", "Pase de lista por grupo y coach/coordinador.", "Modo offline ligero o cache local si sedes tienen mala red."],
        ["Cobranza", "Cargos, pagos parciales, CLABE simulada, tarjeta/link simulado, efectivo aceptado.", "Proveedor real SPEI/terminal y reglas de mora automatizadas."],
        ["Gastos", "Captura y aprobacion basica.", "Flujo formal de solicitud, evidencia obligatoria y deduplicacion."],
        ["Contador", "Portal contable, grafica X/Y y export Excel.", "Formato final para estado de resultados semanal."],
        ["Coach", "Equipo, formacion, banca, asistencia y horas.", "Conectar horas con aprobacion de nomina."],
        ["Auditoria", "Modelo existe; uso todavia parcial.", "Middleware o signals para registrar cambios criticos."],
        ["Despliegue", "No implementado.", "Preparar staging, produccion, backups, dominio y variables."],
    ], [1.2, 2.8, 2.4])
    doc.add_heading("3. Gantt propuesto", level=1)
    doc.add_picture(str(gantt_png), width=Inches(6.5))
    doc.add_heading("4. Sprint 1", level=1)
    bullets(doc, [
        "Objetivo: demo funcional para demostrar control operativo-financiero minimo.",
        "Entregables: login, sedes, alumnos, representantes, asistencia, pagos, gastos, dashboards, contador, coach y documentacion.",
        "Riesgo controlado: no prometer integraciones bancarias reales ni Contpaqi en una semana.",
    ])
    doc.add_heading("5. Sprint 2 propuesto", level=1)
    bullets(doc, [
        "Despliegue en staging y produccion con PostgreSQL, backups y HTTPS.",
        "Importacion masiva de Excel/Gymforce con validaciones de datos sucios.",
        "Reglas de cobranza: vencimientos mensuales, tolerancia de 10 dias, mora propuesta 5%, pagos parciales y topes para jugar.",
        "Proveedor real de pagos: SPEI/CLABE virtual, terminal/link de pago y webhooks.",
        "Auditoria automatica de cambios sensibles y cierre diario de caja/sede.",
        "Primeras metricas deportivas: posicion, categorias, equipos, resultados y relacion coach/alumno.",
    ])
    doc.add_heading("6. Plan de despliegue", level=1)
    add_table(doc, ["Componente", "Decision recomendada"], [
        ["Frontend", "Build Vite servido por Nginx o plataforma tipo Vercel/Netlify si API queda separada."],
        ["Backend", "Django + Gunicorn/Uvicorn detras de Nginx o PaaS."],
        ["Base de datos", "PostgreSQL administrado, no SQLite en produccion."],
        ["Archivos", "S3 compatible para fotos, responsivas, comprobantes e identificaciones."],
        ["Seguridad", "HTTPS, variables de entorno, CORS cerrado, backups, roles y politicas de acceso."],
        ["Observabilidad", "Logs de aplicacion, errores, metricas de API y alertas de fallas de pago/webhook."],
        ["Migracion", "Carga controlada con plantilla Excel, validacion y corrida piloto por sede."],
    ], [1.5, 4.9])
    doc.add_heading("7. Riesgos principales", level=1)
    bullets(doc, [
        "Datos personales de menores: requiere permisos, almacenamiento seguro y minimo acceso.",
        "Efectivo: el sistema ayuda a auditar, pero requiere corte fisico y responsable.",
        "Reglas ambiguas: precios, mora, descuentos y topes deben cerrarse antes de automatizar.",
        "Integracion bancaria: depende de proveedor, contrato, webhooks y pruebas externas.",
        "Adopcion: coordinadores y cajeros necesitan flujos rapidos, no pantallas administrativas densas.",
    ])
    path = DOCS / "Futsi_Gobernanza_Roadmap.docx"
    doc.save(path)
    return path


def pdf_from_story(path: Path, title: str, story: list) -> None:
    doc = SimpleDocTemplate(str(path), pagesize=letter, rightMargin=0.65 * inch, leftMargin=0.65 * inch, topMargin=0.65 * inch, bottomMargin=0.65 * inch)
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10))
    styles["Title"].textColor = colors.HexColor("#0f172a")
    styles["Heading1"].textColor = colors.HexColor("#14532d")
    doc.build(story)


def make_pdf_documents(erd_png: Path, gantt_png: Path) -> None:
    styles = getSampleStyleSheet()
    req_story = [
        Paragraph("Especificacion de requerimientos - Futsi Mini ERP", styles["Title"]),
        Paragraph("Version Sprint 1 actualizada al 26 de mayo de 2026.", styles["BodyText"]),
        Spacer(1, 0.15 * inch),
        Paragraph("Alcance", styles["Heading1"]),
        Paragraph("Mini ERP operativo-financiero para controlar sedes, alumnos, representantes, equipos, asistencia, pagos, descuentos, gastos, dashboards y trazabilidad. No sustituye Contpaqi ni resuelve fiscalidad completa en Sprint 1.", styles["BodyText"]),
        Paragraph("Requerimientos funcionales principales", styles["Heading1"]),
        ListFlowable([ListItem(Paragraph(text, styles["BodyText"])) for text in [
            "Login por roles: admin, direccion, contador, coordinador, cajero, coach y representante.",
            "Catalogo de sedes con coordenadas y mapa real.",
            "Alumnos con foto, datos medicos, responsiva, uniforme, descuentos, pausas y representante.",
            "Pase de lista por sede/grupo, con indicador de adeudo y cierre.",
            "Cobranza: cargos, pagos parciales, efectivo con aceptacion, transferencia CLABE simulada, terminal y link.",
            "Gastos: captura, estatus y aprobacion.",
            "Portal contador: grafica X/Y y exportacion Excel.",
            "Portal coach: formacion 4-3-3, banca, asistencia y horas.",
        ]], bulletType="bullet"),
        Paragraph("Modelo de datos", styles["Heading1"]),
        RLImage(str(erd_png), width=7.0 * inch, height=5.13 * inch),
        PageBreak(),
        Paragraph("Criterios de aceptacion", styles["Heading1"]),
        ListFlowable([ListItem(Paragraph(text, styles["BodyText"])) for text in [
            "El contador exporta Excel con resumen, pagos, cargos, gastos, descuentos y cruces de asistencia con adeudo.",
            "El cajero no ve informacion ejecutiva completa y solo cobra su sede.",
            "El coach solo ve su grupo asignado.",
            "El representante solo ve sus alumnos y pagos.",
            "La demo permite explicar el 80/20: asistencia + cobranza + gastos + descuentos + auditoria.",
        ]], bulletType="bullet"),
    ]
    pdf_from_story(DOCS / "Futsi_Especificacion_Requerimientos.pdf", "req", req_story)

    gov_table = Table([
        ["Area", "Estado", "Siguiente paso"],
        ["Roles", "Implementado", "Afinar permisos con responsables reales"],
        ["Contador", "Portal, grafica, Excel", "Formato final de estado de resultados"],
        ["Pagos", "Simulado", "Proveedor real SPEI/terminal"],
        ["Auditoria", "Parcial", "Signals/middleware de cambios criticos"],
        ["Despliegue", "Pendiente", "Staging, produccion, backups, HTTPS"],
    ], colWidths=[1.35 * inch, 2.0 * inch, 3.3 * inch])
    gov_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#dcfce7")),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#cbd5e1")),
        ("FONT", (0, 0), (-1, -1), "Helvetica", 8),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
    ]))
    gov_story = [
        Paragraph("Gobernanza, roadmap y control de ejecucion", styles["Title"]),
        Paragraph("Sprint 1 y Sprint 2 propuesto para Futsi Mini ERP.", styles["BodyText"]),
        Paragraph("Principio", styles["Heading1"]),
        Paragraph("Primero control y trazabilidad; despues automatizacion, integraciones y analitica avanzada.", styles["BodyText"]),
        Paragraph("Checklist", styles["Heading1"]),
        gov_table,
        Spacer(1, 0.2 * inch),
        Paragraph("Gantt", styles["Heading1"]),
        RLImage(str(gantt_png), width=7.0 * inch, height=3.62 * inch),
        PageBreak(),
        Paragraph("Despliegue y proximos pasos", styles["Heading1"]),
        ListFlowable([ListItem(Paragraph(text, styles["BodyText"])) for text in [
            "Preparar ambiente staging con PostgreSQL, variables de entorno, HTTPS y CORS cerrado.",
            "Configurar almacenamiento S3 compatible para fotos, responsivas y comprobantes.",
            "Definir proveedor de pagos para CLABE virtual, SPEI, terminal y links.",
            "Implementar auditoria automatica para pagos, descuentos, gastos, asistencia y cierres.",
            "Migrar datos reales desde Excel/Gymforce y hacer piloto por sede.",
            "Cerrar reglas de mora, topes para jugar, reembolsos, pagos parciales y descuentos.",
        ]], bulletType="bullet"),
    ]
    pdf_from_story(DOCS / "Futsi_Gobernanza_Roadmap.pdf", "gov", gov_story)


def main() -> None:
    erd_png = create_erd_files()
    gantt_png = create_gantt_image()
    requirements_doc(erd_png)
    governance_doc(gantt_png)
    make_pdf_documents(erd_png, gantt_png)
    print("Deliverables created in", DOCS)


if __name__ == "__main__":
    main()
