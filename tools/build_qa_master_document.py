from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import (
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
DOCX_PATH = DOCS / "Futsi_QA_Master_Sprint4.docx"
PDF_PATH = DOCS / "Futsi_QA_Master_Sprint4.pdf"
MD_PATH = DOCS / "Futsi_QA_Master_Sprint4.md"


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(str(text))
    run.bold = bold
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_docx_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(header_cells[idx], header, bold=True)
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    doc.add_paragraph()
    return table


def add_bullet(doc, text):
    paragraph = doc.add_paragraph(style="List Bullet")
    paragraph.add_run(text)


def build_markdown():
    content = f"""# Documento Maestro de QA - Sprint 4

Fecha de corte: {date.today().isoformat()}

## Resumen ejecutivo

El Sprint 4 agrega una capa formal de QA automatizado sobre Futsi Mini ERP. La suite queda dividida en pruebas de API/backend con Pytest, pruebas end-to-end web con Selenium, build de frontend, evidencia automatica de fallos y analisis estatico con SonarQube/SonarCloud. El objetivo es detectar regresiones en permisos, cobranza, facturacion, carga historica, navegacion por roles y flujos moviles antes de desplegar.

## Resultados actuales

| Tipo | Herramienta | Resultado | Cobertura / evidencia |
|---|---|---|---|
| Backend/API | Pytest + pytest-django | 27 pruebas esperadas despues de agregar rol dev | Coverage minimo CI: 60%; ultima corrida previa: 83.72% |
| E2E Web | Selenium + Chrome headless | 6 pruebas pasan | Captura PNG/HTML ante fallo en qa/artifacts |
| Frontend | Vite build | Build exitoso | dist generado localmente |
| Seguridad basica | Pytest security inputs | Incluido en backend | Login SQL-like, montos negativos, permisos cruzados |
| Analisis estatico | SonarQube/SonarCloud | Configurado | sonar-project.properties + job CI condicionado a SONAR_TOKEN |

## Clasificacion de pruebas

### Pruebas backend/API

- Autenticacion y permisos: roles admin, dev, contador, cajero, coach y tutor.
- Cobranza: pagos, efectivo con confirmacion, transferencia simulada, restricciones por sede.
- Reporte contable: exportacion XLSX valida y permisos de contador.
- Facturacion simulada: generacion de PDF/XML/UUID para ingresos y egresos.
- Historico Excel: preview, commit, firma, password y bloqueo por rol no autorizado.
- Robustez y seguridad: payloads tipo SQL injection, textos raros, montos negativos, IDs invalidos.

### Pruebas end-to-end web

- Login invalido no accede al sistema.
- Admin/dev: dashboard, alumnos, historico, tema oscuro y menu movil.
- Contador: exportacion contable y facturas.
- Cajero: controles de pago sin acceso a panel admin.
- Coach: asistencia, camara y registro de horas.
- Tutor: alumnos vinculados, perfil y facturas.

### Pruebas de integracion y build

- Django migrations + seed demo reproducible.
- Vite build de React.
- Selenium levanta Django y Vite con base SQLite aislada.

## Estrategia SonarQube

Se agrego `sonar-project.properties` para centralizar fuentes, pruebas, exclusiones y reportes de cobertura. El job `sonarqube` en GitHub Actions corre Pytest con coverage y ejecuta `SonarSource/sonarqube-scan-action` cuando existe `SONAR_TOKEN`.

Configuracion requerida en GitHub:

- Secret: `SONAR_TOKEN`.
- Variable opcional: `SONAR_HOST_URL`.
- Para SonarCloud usar `https://sonarcloud.io`.
- Para SonarQube Server usar la URL interna/publica del servidor.

Politica recomendada:

- Quality Gate obligatorio antes de deploy productivo.
- Minimo inicial backend: 60% coverage.
- Sin nuevos bugs criticos ni vulnerabilidades criticas.
- Excluir generados: `front/android`, `dist`, migraciones, artifacts y docs binarios.
- Agregar coverage frontend con Vitest en un sprint posterior.

## Rol tecnico Dev App

Se agrega el rol `dev` y el usuario demo `dev/dev12345`. Tiene permisos equivalentes a admin para QA, soporte y diagnostico, pero debe tratarse como cuenta tecnica, no operativa.

Reglas:

- Puede ver usuarios, sedes, alumnos, historicos, reportes y flujos admin.
- Debe usarse para pruebas internas, debugging y validacion en ambientes no productivos.
- En produccion debe tener MFA, password fuerte, auditoria y acceso temporal.
- No debe compartirse entre personas; cada desarrollador deberia tener cuenta propia cuando haya identidad real.

## Matriz de aceptacion

| Area | Criterio | Estado |
|---|---|---|
| Backend pytest | Suite completa pasa | Implementado |
| Coverage | Minimo 60% | Implementado |
| Selenium | Roles principales cubiertos | Implementado |
| Artifacts | Screenshot/HTML ante fallo | Implementado |
| SonarQube | Config y CI preparado | Implementado |
| Dev App | Rol y usuario tecnico | Implementado |
| DeepFace | Fuera de CI; demo local | Decidido |
| Android nativo QA | Pendiente automatizar | Futuro |

## Riesgos y siguientes pasos

- SonarQube necesita `SONAR_TOKEN` real para escanear en CI.
- Falta agregar pruebas frontend unitarias con Vitest para generar `front/coverage/lcov.info`.
- Selenium cubre smoke e2e; aun falta ampliar casos destructivos de formularios y doble click.
- Android queda cubierto por responsive web y smoke manual; Appium puede entrar despues.
- El rol dev debe endurecerse antes de produccion: MFA, expiracion, audit log y principio de menor privilegio.
"""
    MD_PATH.write_text(content, encoding="utf-8")
    return content


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = doc.styles
    styles["Normal"].font.name = "Calibri"
    styles["Normal"].font.size = Pt(11)
    styles["Heading 1"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 2"].font.color.rgb = RGBColor(46, 116, 181)
    styles["Heading 3"].font.color.rgb = RGBColor(31, 77, 120)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Documento Maestro de QA - Sprint 4")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run(f"Futsi Mini ERP | Fecha de corte: {date.today().isoformat()}").italic = True

    doc.add_heading("1. Resumen ejecutivo", level=1)
    doc.add_paragraph(
        "El Sprint 4 agrega una capa formal de QA automatizado sobre Futsi Mini ERP. "
        "La suite queda dividida en pruebas de API/backend con Pytest, pruebas end-to-end web con Selenium, "
        "build de frontend, evidencia automatica de fallos y analisis estatico con SonarQube/SonarCloud."
    )

    doc.add_heading("2. Resultados actuales", level=1)
    add_docx_table(
        doc,
        ["Tipo", "Herramienta", "Resultado", "Evidencia"],
        [
            ["Backend/API", "Pytest", "27 pruebas esperadas con rol dev", "Coverage minimo CI 60%; ultima corrida previa 83.72%"],
            ["E2E Web", "Selenium", "6 pruebas pasan", "PNG/HTML ante fallo"],
            ["Frontend", "Vite", "Build exitoso", "dist generado localmente"],
            ["Seguridad basica", "Pytest", "Incluido", "SQL-like, permisos, montos negativos"],
            ["Analisis estatico", "SonarQube", "Configurado", "sonar-project.properties + CI"],
        ],
    )

    doc.add_heading("3. Clasificacion de pruebas", level=1)
    doc.add_heading("Backend/API", level=2)
    for item in [
        "Autenticacion y permisos por rol: admin, dev, contador, cajero, coach y tutor.",
        "Cobranza: pagos, efectivo con confirmacion, transferencia simulada y restricciones por sede.",
        "Reporte contable: exportacion XLSX valida y permisos de contador.",
        "Facturacion simulada: PDF/XML/UUID para ingresos y egresos.",
        "Historico Excel: preview, commit, firma, password y bloqueo por rol.",
        "Robustez: SQL-like payloads, textos raros, montos negativos e IDs invalidos.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("End-to-end web", level=2)
    for item in [
        "Login invalido no accede al sistema.",
        "Admin/dev: dashboard, alumnos, historico, tema oscuro y menu movil.",
        "Contador: exportacion contable y facturas.",
        "Cajero: controles de pago sin acceso admin.",
        "Coach: asistencia, camara y registro de horas.",
        "Tutor: alumnos vinculados, perfil y facturas.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("4. Estrategia SonarQube", level=1)
    doc.add_paragraph(
        "Se agrego sonar-project.properties para centralizar fuentes, pruebas, exclusiones y reportes. "
        "El job sonarqube corre Pytest con coverage y ejecuta SonarSource/sonarqube-scan-action cuando existe SONAR_TOKEN."
    )
    add_docx_table(
        doc,
        ["Elemento", "Decision"],
        [
            ["Secret requerido", "SONAR_TOKEN"],
            ["Variable opcional", "SONAR_HOST_URL"],
            ["SonarCloud", "https://sonarcloud.io"],
            ["Quality Gate", "Obligatorio antes de deploy productivo"],
            ["Exclusiones", "android generado, dist, migraciones, artifacts y docs binarios"],
        ],
    )

    doc.add_heading("5. Rol tecnico Dev App", level=1)
    doc.add_paragraph(
        "Se agrega el rol dev y el usuario demo dev/dev12345. Tiene permisos equivalentes a admin para QA, soporte y diagnostico, "
        "pero debe tratarse como cuenta tecnica, no operativa."
    )
    for item in [
        "Puede ver usuarios, sedes, alumnos, historicos, reportes y flujos admin.",
        "Debe usarse para pruebas internas y validacion en ambientes no productivos.",
        "En produccion requiere MFA, password fuerte, auditoria y acceso temporal.",
        "No debe compartirse; cada desarrollador deberia tener cuenta individual.",
    ]:
        add_bullet(doc, item)

    doc.add_heading("6. Matriz de aceptacion", level=1)
    add_docx_table(
        doc,
        ["Area", "Criterio", "Estado"],
        [
            ["Backend pytest", "Suite completa pasa", "Implementado"],
            ["Coverage", "Minimo 60%", "Implementado"],
            ["Selenium", "Roles principales cubiertos", "Implementado"],
            ["Artifacts", "Screenshot/HTML ante fallo", "Implementado"],
            ["SonarQube", "Config y CI preparado", "Implementado"],
            ["Dev App", "Rol y usuario tecnico", "Implementado"],
            ["DeepFace", "Fuera de CI; demo local", "Decidido"],
            ["Android nativo QA", "Automatizacion futura", "Pendiente"],
        ],
    )

    doc.add_heading("7. Riesgos y siguientes pasos", level=1)
    for item in [
        "Configurar SONAR_TOKEN real para activar el scan en CI.",
        "Agregar pruebas frontend unitarias con Vitest para producir front/coverage/lcov.info.",
        "Ampliar Selenium con casos destructivos de formularios y doble click.",
        "Evaluar Appium para automatizacion Android nativa.",
        "Endurecer el rol dev antes de produccion con MFA, expiracion y auditoria.",
    ]:
        add_bullet(doc, item)

    doc.save(DOCX_PATH)


def paragraph(text, style):
    return Paragraph(text, style)


def build_pdf():
    styles = getSampleStyleSheet()
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10))
    styles["Title"].textColor = colors.HexColor("#0B2545")
    styles["Heading1"].textColor = colors.HexColor("#2E74B5")
    styles["Heading2"].textColor = colors.HexColor("#2E74B5")

    pdf = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        rightMargin=inch,
        leftMargin=inch,
        topMargin=inch,
        bottomMargin=inch,
    )
    story = [
        paragraph("Documento Maestro de QA - Sprint 4", styles["Title"]),
        paragraph(f"Futsi Mini ERP | Fecha de corte: {date.today().isoformat()}", styles["BodyText"]),
        Spacer(1, 0.18 * inch),
        paragraph("Resumen ejecutivo", styles["Heading1"]),
        paragraph(
            "El Sprint 4 agrega QA automatizado con Pytest, Selenium, evidencia de fallos y SonarQube/SonarCloud.",
            styles["BodyText"],
        ),
        paragraph("Resultados actuales", styles["Heading1"]),
    ]
    table_data = [
        ["Tipo", "Herramienta", "Resultado"],
        ["Backend/API", "Pytest", "27 pruebas esperadas; coverage minimo 60%"],
        ["E2E Web", "Selenium", "6 pruebas pasan; artifacts ante fallo"],
        ["Frontend", "Vite", "Build exitoso"],
        ["SonarQube", "Scan action", "Config preparado y condicionado a SONAR_TOKEN"],
    ]
    table = Table(table_data, colWidths=[1.3 * inch, 1.5 * inch, 3.4 * inch])
    table.setStyle(
        TableStyle(
            [
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#F2F4F7")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#D0D5DD")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
                ("FONTSIZE", (0, 0), (-1, -1), 8),
                ("LEFTPADDING", (0, 0), (-1, -1), 6),
                ("RIGHTPADDING", (0, 0), (-1, -1), 6),
            ]
        )
    )
    story.extend([table, Spacer(1, 0.18 * inch)])
    for heading, items in [
        (
            "Clasificacion",
            [
                "Backend/API: permisos, pagos, reportes, facturas, historico Excel y seguridad basica.",
                "E2E: login, admin/dev, contador, cajero, coach y tutor.",
                "Build/integracion: Django, seed demo, Vite y SQLite aislado.",
            ],
        ),
        (
            "Estrategia SonarQube",
            [
                "Usar SONAR_TOKEN y SONAR_HOST_URL en GitHub.",
                "Quality Gate obligatorio antes de deploy productivo.",
                "Agregar Vitest para coverage frontend en siguiente fase.",
            ],
        ),
        (
            "Rol Dev App",
            [
                "Usuario demo dev/dev12345 con permisos equivalentes a admin para QA.",
                "En produccion: MFA, auditoria, expiracion y cuentas individuales.",
            ],
        ),
    ]:
        story.append(paragraph(heading, styles["Heading1"]))
        for item in items:
            story.append(paragraph(f"- {item}", styles["BodyText"]))
    pdf.build(story)


def main():
    DOCS.mkdir(exist_ok=True)
    build_markdown()
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
    print(MD_PATH)


if __name__ == "__main__":
    main()
