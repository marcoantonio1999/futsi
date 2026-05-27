from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    KeepTogether,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
DOCS.mkdir(exist_ok=True)


ACCENT = colors.HexColor("#14532d")
BORDER = colors.HexColor("#cbd5e1")
HEADER_FILL = colors.HexColor("#dcfce7")
INK = colors.HexColor("#0f172a")
MUTED = colors.HexColor("#52525b")


def styles():
    base = getSampleStyleSheet()
    base["Title"].fontName = "Helvetica-Bold"
    base["Title"].fontSize = 22
    base["Title"].leading = 26
    base["Title"].textColor = INK
    base["Heading1"].fontName = "Helvetica-Bold"
    base["Heading1"].fontSize = 15
    base["Heading1"].leading = 18
    base["Heading1"].spaceBefore = 14
    base["Heading1"].spaceAfter = 7
    base["Heading1"].textColor = ACCENT
    base["Heading2"].fontName = "Helvetica-Bold"
    base["Heading2"].fontSize = 12
    base["Heading2"].leading = 15
    base["Heading2"].spaceBefore = 10
    base["Heading2"].spaceAfter = 5
    base["Heading2"].textColor = INK
    base["BodyText"].fontName = "Helvetica"
    base["BodyText"].fontSize = 9.2
    base["BodyText"].leading = 12
    base["BodyText"].spaceAfter = 5
    base.add(ParagraphStyle(name="Small", parent=base["BodyText"], fontSize=8, leading=10, textColor=MUTED))
    base.add(ParagraphStyle(name="Callout", parent=base["BodyText"], backColor=colors.HexColor("#f0fdf4"), borderColor=colors.HexColor("#bbf7d0"), borderWidth=0.6, borderPadding=7, spaceBefore=6, spaceAfter=8))
    return base


def p(text: str, style):
    return Paragraph(text, style)


def bullet_list(items: list[str], style):
    return ListFlowable([ListItem(p(item, style)) for item in items], bulletType="bullet", start="circle", leftIndent=16)


def numbered_list(items: list[str], style):
    return ListFlowable([ListItem(p(item, style)) for item in items], bulletType="1", leftIndent=18)


def table(data: list[list[str]], widths: list[float], repeat: bool = True):
    converted = [[p(str(cell), styles()["Small"]) for cell in row] for row in data]
    t = Table(converted, colWidths=[w * inch for w in widths], repeatRows=1 if repeat else 0)
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HEADER_FILL),
        ("TEXTCOLOR", (0, 0), (-1, 0), INK),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 8),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 8),
        ("GRID", (0, 0), (-1, -1), 0.35, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return t


def pdf(path: Path, story: list):
    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.65 * inch, 0.38 * inch, "Futsi Mini ERP - documento de trabajo")
        canvas.drawRightString(7.85 * inch, 0.38 * inch, f"Pagina {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def docx_style(doc: Document):
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    for name, size, color in [("Heading 1", 15, "14532D"), ("Heading 2", 12, "0F172A"), ("Heading 3", 10.5, "27272A")]:
        s = doc.styles[name]
        s.font.name = "Arial"
        s.font.size = Pt(size)
        s.font.bold = True
        s.font.color.rgb = RGBColor.from_string(color)


def docx_title(doc: Document, title: str, subtitle: str):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = para.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0F172A")
    sub = doc.add_paragraph(subtitle)
    sub.runs[0].font.name = "Arial"
    sub.runs[0].font.size = Pt(10)
    sub.runs[0].font.color.rgb = RGBColor.from_string("52525B")


def docx_table(doc: Document, headers: list[str], rows: list[list[str]]):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"
    for i, header in enumerate(headers):
        t.rows[0].cells[i].text = header
        for run in t.rows[0].cells[i].paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            for para in cells[i].paragraphs:
                for run in para.runs:
                    run.font.size = Pt(8)


def add_docx_bullets(doc: Document, items: list[str]):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def srs_story():
    s = styles()
    erd = ASSETS / "db_schema_erd.png"
    story = [
        p("Especificacion de requerimientos de software", s["Title"]),
        p("Futsi Mini ERP operativo-financiero para academias y torneos de futbol", s["BodyText"]),
        p("Version: Sprint 1 + planeacion Sprint 2. Fecha: 26 de mayo de 2026.", s["Small"]),
        Spacer(1, 0.12 * inch),
        p("Resumen ejecutivo", s["Heading1"]),
        p("El proyecto busca reducir fugas de ingreso, errores operativos y falta de trazabilidad en una empresa que opera academias infantiles y torneos de futbol para adultos en varias sedes de CDMX. La solucion no pretende sustituir Contpaqi ni convertirse en un ERP fiscal completo; su primer valor es crear una fuente confiable para cruzar asistencia, cobranza, descuentos, gastos y cierres por sede.", s["Callout"]),
        p("Control de documento", s["Heading1"]),
        table([
            ["Campo", "Valor"],
            ["Nombre", "Futsi Mini ERP - Especificacion de requerimientos"],
            ["Audiencia", "Direccion, contador, operaciones, desarrollo y evaluadores del Sprint"],
            ["Stack actual", "React/Vite/Tailwind, Django REST Framework, SQLite demo, PostgreSQL recomendado para produccion"],
            ["Estado", "Documento formal de trabajo; debe validarse con usuarios finales antes de produccion"],
        ], [1.6, 5.2], False),
        p("Indice funcional", s["Heading1"]),
        numbered_list([
            "Contexto y problema operativo.",
            "Objetivos, alcance, exclusiones y supuestos.",
            "Actores, roles y permisos.",
            "Requerimientos funcionales por modulo.",
            "Reglas de negocio y flujos principales.",
            "Requerimientos no funcionales.",
            "Modelo de datos, reportes, riesgos, despliegue y criterios de aceptacion.",
        ], s["BodyText"]),
        PageBreak(),
        p("1. Contexto de negocio", s["Heading1"]),
        p("La empresa opera academias de futbol para ninos de 4 a 17 anos y torneos de futbol para adultos. La operacion ocurre por las tardes/noches entre semana y durante fines de semana. Los pagos se reciben por efectivo, transferencia y tarjeta; una parte importante se registra en ventanilla y otra llega por comprobantes enviados por WhatsApp.", s["BodyText"]),
        bullet_list([
            "Aproximadamente 400 equipos y 1000 alumnos de academia como escala objetivo.",
            "Los coordinadores tienen perfil operativo/deportivo; por tanto la interfaz debe ser simple.",
            "El cajero/auxiliar administrativo captura pagos y usa terminal para tarjeta.",
            "El contador necesita estados de resultados, pagos/no pagos y detalle exportable.",
            "Una jornada equivale a una semana de juegos; puede existir doble jornada.",
            "Los torneos duran 12 jornadas mas liguilla.",
        ], s["BodyText"]),
        p("2. Problema actual", s["Heading1"]),
        bullet_list([
            "Excel permite cambios posteriores al corte, dificulta saber quien edito y puede contener formulas rotas.",
            "Gymforce no modela bien academias y ligas; se usa parcialmente y luego se transcribe a Excel.",
            "Los descuentos pueden aplicarse sin evidencia suficiente y deben cruzarse contra asistencia.",
            "El efectivo y las transferencias no conciliadas son fuentes de riesgo.",
            "Los gastos pueden duplicarse, inflarse o asignarse a una sede incorrecta.",
            "La informacion se presenta tarde y con poco historial auditable.",
        ], s["BodyText"]),
        p("3. Objetivos", s["Heading1"]),
        table([
            ["Objetivo", "Descripcion", "Medicion"],
            ["Control de ingresos", "Identificar quien debia pagar, quien pago y quien recibio el pago.", "Pagos y cargos conciliables por alumno/equipo/sede."],
            ["Trazabilidad", "Registrar usuario, fecha, accion y cambios relevantes.", "Bitacora para pagos, descuentos, gastos, asistencia y cierres."],
            ["Operacion simple", "Permitir que cancha y ventanilla operen rapido.", "Pase de lista y cobro con pocos pasos."],
            ["Visibilidad contable", "Dar al contador reportes y Excel detallado.", "Exportacion y dashboard financiero."],
            ["Base escalable", "Preparar Sprint 2 sin sobreingenieria.", "Modelo de datos e integraciones claras."],
        ], [1.4, 3.5, 1.9]),
        p("4. Alcance de Sprint 1", s["Heading1"]),
        bullet_list([
            "Login con roles y permisos.",
            "Catalogos base: sedes, representantes, alumnos, equipos/torneos base.",
            "Pase de lista por sede/grupo y portal coach.",
            "Cobranza: cargos, pagos parciales y simulacion de transferencia CLABE, efectivo confirmado, terminal y link.",
            "Gastos operativos con estados de aprobacion.",
            "Dashboard admin/direccion y portal contable con exportacion Excel.",
            "Documentacion, datos demo y pruebas automatizadas.",
        ], s["BodyText"]),
        p("5. Fuera de alcance Sprint 1", s["Heading1"]),
        bullet_list([
            "Reemplazar Contpaqi, facturacion fiscal/CFDI o contabilidad fiscal completa.",
            "Integracion real con banco, SPEI, terminal fisica o proveedor de pagos.",
            "Camaras, reconocimiento facial, kiosko ATM o app movil nativa.",
            "Modo offline completo.",
            "Inventario avanzado de uniformes.",
            "Estadisticas deportivas avanzadas completas.",
        ], s["BodyText"]),
        p("6. Actores y permisos", s["Heading1"]),
        table([
            ["Rol", "Necesidad", "Permisos esperados"],
            ["Admin", "Configurar y operar demo completa.", "CRUD de sedes, usuarios, alumnos, pagos, gastos, descuentos, reportes."],
            ["Direccion", "Ver salud operativa y financiera.", "Dashboard consolidado, mapa, metricas y alertas."],
            ["Contador", "Preparar estado de resultados y revisar pagos/no pagos.", "Consulta contable, grafica X/Y, exportacion Excel; sin configurar usuarios."],
            ["Coordinador", "Gerente de sede y cancha.", "Asistencia, autorizaciones operativas, seguimiento de sede."],
            ["Cajero", "Cobrar en ventanilla.", "Procesar pagos de su sede; sin ver datos ejecutivos."],
            ["Coach", "Entrenar y pasar lista.", "Ver su grupo/equipo, registrar horas, asistencia y alertas."],
            ["Representante", "Pagar y validar informacion de sus hijos.", "Ver adeudos, CLABE, links, pagos y confirmar efectivo."],
        ], [1.1, 2.4, 3.3]),
        PageBreak(),
        p("7. Requerimientos funcionales detallados", s["Heading1"]),
        p("7.1 Usuarios, autenticacion y alcance", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-001", "El sistema debe permitir login por usuario y password.", "Alta", "Implementado"],
            ["RF-002", "El sistema debe distinguir roles: admin, direccion, contador, coordinador, cajero, coach y representante.", "Alta", "Implementado"],
            ["RF-003", "El cajero solo debe ver datos necesarios para cobrar en su sede.", "Alta", "Implementado"],
            ["RF-004", "El coach solo debe ver su grupo/equipo asignado.", "Alta", "Implementado"],
            ["RF-005", "El representante solo debe ver alumnos y pagos asociados a su perfil.", "Alta", "Implementado"],
            ["RF-006", "El contador debe tener portal contable diferenciado del admin.", "Alta", "Implementado"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.2 Sedes y configuracion", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-010", "Registrar sedes con nombre, codigo, direccion, coordenadas y ventana de edicion.", "Alta", "Implementado"],
            ["RF-011", "Mostrar mapa real con sedes y datos operativos.", "Media", "Implementado"],
            ["RF-012", "Permitir configuracion futura de precios por sede.", "Alta", "Pendiente Sprint 2"],
            ["RF-013", "Permitir reglas de cierre por sede.", "Alta", "Parcial"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.3 Alumnos y representantes", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-020", "Registrar alumno con nombre, fecha nacimiento, sede, categoria, grupo, estado, foto y responsiva.", "Alta", "Implementado"],
            ["RF-021", "Registrar informacion medica, contacto de emergencia y telefono.", "Alta", "Implementado"],
            ["RF-022", "Registrar uniforme pendiente, pagado o entregado.", "Media", "Implementado"],
            ["RF-023", "Registrar pausas autorizadas con fechas y motivo.", "Alta", "Implementado"],
            ["RF-024", "Asociar uno o varios alumnos a un representante.", "Alta", "Implementado"],
            ["RF-025", "Generar CLABE virtual simulada por representante.", "Alta", "Implementado"],
            ["RF-026", "Importar alumnos desde Excel/Gymforce.", "Alta", "Pendiente Sprint 2"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.4 Asistencia", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-030", "Crear sesiones de asistencia por sede, fecha, hora y grupo.", "Alta", "Implementado"],
            ["RF-031", "Marcar asistio, falto o justificada.", "Alta", "Implementado"],
            ["RF-032", "Mostrar adeudo al momento de pasar lista.", "Alta", "Implementado"],
            ["RF-033", "Guardar si habia adeudo al capturar asistencia.", "Alta", "Implementado"],
            ["RF-034", "Cerrar sesion de asistencia para evitar edicion libre.", "Alta", "Implementado base"],
            ["RF-035", "Soportar modo offline parcial.", "Media", "Pendiente"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.5 Cobranza y pagos", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-040", "Crear cargos por mensualidad, torneo, jornada, uniforme, sancion u otro concepto.", "Alta", "Implementado"],
            ["RF-041", "Aceptar pagos parciales y recalcular saldo.", "Alta", "Implementado"],
            ["RF-042", "Simular transferencia con CLABE virtual y estado en proceso hasta webhook.", "Alta", "Implementado demo"],
            ["RF-043", "Simular pago en efectivo con aceptacion del representante.", "Alta", "Implementado demo"],
            ["RF-044", "Simular pago con terminal o link de tarjeta.", "Alta", "Implementado demo"],
            ["RF-045", "Expirar pagos de transferencia despues de 72 horas si no se confirman.", "Media", "Implementado accion demo"],
            ["RF-046", "Aplicar mora de 5% despues de 10 dias.", "Alta", "Pendiente Sprint 2"],
            ["RF-047", "Bloquear juego/entrenamiento por niveles de aviso.", "Alta", "Pendiente Sprint 2"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.6 Descuentos", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-050", "Solicitar descuentos con motivo y monto.", "Alta", "Implementado"],
            ["RF-051", "Aprobar o rechazar descuentos.", "Alta", "Implementado"],
            ["RF-052", "Cruzar descuentos contra asistencia para detectar riesgo.", "Alta", "Parcial"],
            ["RF-053", "Configurar descuentos homogeneos: hermano 15%, referido y promociones.", "Alta", "Pendiente Sprint 2"],
            ["RF-054", "Adjuntar evidencia obligatoria segun tipo de descuento.", "Media", "Pendiente"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.7 Gastos y nomina operativa", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-060", "Capturar gastos por sede, categoria, proveedor, monto y fecha.", "Alta", "Implementado"],
            ["RF-061", "Aprobar o rechazar gastos.", "Alta", "Implementado"],
            ["RF-062", "Registrar horas de coach con tarifa snapshot.", "Media", "Implementado"],
            ["RF-063", "Cruzar horas de coach contra asistencia/grupo.", "Media", "Pendiente Sprint 2"],
            ["RF-064", "Detectar gastos duplicados o sin comprobante.", "Alta", "Pendiente Sprint 2"],
        ], [0.65, 4.2, 0.8, 1.0]),
        p("7.8 Contador, dashboard y reportes", s["Heading2"]),
        table([
            ["ID", "Requerimiento", "Prioridad", "Estado"],
            ["RF-070", "Mostrar ingresos, egresos, utilidad y saldo por cobrar.", "Alta", "Implementado"],
            ["RF-071", "Mostrar grafica con eje X/Y de ingresos, egresos y utilidad.", "Alta", "Implementado"],
            ["RF-072", "Exportar archivo Excel con detalle contable.", "Alta", "Implementado"],
            ["RF-073", "Mostrar alertas de alumnos que asistieron con adeudo.", "Alta", "Implementado"],
            ["RF-074", "Estado de resultados semanal/mensual final.", "Alta", "Pendiente formato empresa"],
        ], [0.65, 4.2, 0.8, 1.0]),
        PageBreak(),
        p("8. Reglas de negocio", s["Heading1"]),
        table([
            ["Regla", "Descripcion", "Estado"],
            ["Mensualidad", "Academia paga por mes completo; no existe inscripcion inicial segun aclaracion actual.", "Definida"],
            ["Prueba", "Periodo de prueba normal de 2 semanas/2 clases, extensible hasta 4 semanas.", "Pendiente parametrizar"],
            ["Mora academia", "Si no paga en 10 dias, se propone penalizacion de 5%.", "Pendiente Sprint 2"],
            ["Torneos", "Puede cobrarse por torneo completo o por jornada semanal.", "Modelo base"],
            ["Jornada", "Una jornada equivale a una semana de juegos; puede existir doble jornada.", "Definida"],
            ["Torneo", "Duracion: 12 jornadas mas liguilla.", "Pendiente modelado completo"],
            ["Equipo con adeudo", "Puede jugar con decision del coordinador, pero debe acumular mora y tener tope.", "Pendiente Sprint 2"],
            ["Avisos a padres", "Primer aviso a papas, segundo al nino, tercero no juega partidos, cuarto no entrena.", "Pendiente Sprint 2"],
            ["Transferencia", "Debe usarse CLABE virtual por cliente y webhook para confirmacion automatica.", "Simulado"],
            ["Efectivo", "Cajero registra solicitud y representante acepta para dejar evidencia bilateral.", "Simulado"],
            ["Tarjeta", "Terminal o link debe confirmar automaticamente.", "Simulado"],
        ], [1.35, 4.2, 1.25]),
        p("9. Requerimientos no funcionales", s["Heading1"]),
        table([
            ["Categoria", "Requerimiento"],
            ["Usabilidad", "Interfaz minimalista, clara, responsive y usable desde tablet/celular en cancha."],
            ["Rendimiento", "Debe soportar escala objetivo inicial de 1000 alumnos y 400 equipos sin degradacion perceptible."],
            ["Seguridad", "Autenticacion por token, CORS restringido, HTTPS en produccion, secretos fuera del repositorio."],
            ["Privacidad", "Datos de menores, fotos, responsivas, identificaciones y salud deben tener acceso minimo necesario."],
            ["Auditoria", "Cambios sensibles deben registrar actor, fecha, entidad, valores previos y nuevos."],
            ["Disponibilidad", "Backups de base de datos y archivos; monitoreo de backend y proveedor de pagos."],
            ["Mantenibilidad", "API REST modular, migraciones versionadas y pruebas automatizadas."],
            ["Portabilidad", "Desarrollo local con SQLite; produccion en PostgreSQL."],
        ], [1.4, 5.4]),
        p("10. Modelo de datos", s["Heading1"]),
        p("El modelo se organiza alrededor de sedes, usuarios, alumnos/representantes, torneos/equipos, asistencia, cargos, pagos, descuentos, gastos, horas de coach, cierres y auditoria.", s["BodyText"]),
        Image(str(erd), width=7.0 * inch, height=5.13 * inch),
        PageBreak(),
        p("11. Reportes y exportaciones", s["Heading1"]),
        table([
            ["Reporte", "Usuario", "Contenido"],
            ["Estado de resultados operativo", "Contador / direccion", "Ingresos, egresos, utilidad por sede y consolidado."],
            ["Cobranza abierta", "Contador / cajero / direccion", "Cargos pendientes, parciales y vencidos."],
            ["Asistencia con adeudo", "Direccion / coordinador", "Alumnos/equipos que asistieron/jugaron con adeudo."],
            ["Descuentos", "Admin / contador", "Solicitudes, aprobaciones, motivos, montos y usuario."],
            ["Gastos", "Contador / direccion", "Gastos por sede, categoria, estatus y aprobador."],
            ["Excel contable", "Contador", "Resumen, pagos, cargos, gastos, descuentos y asistencia con adeudo."],
        ], [1.6, 1.6, 3.6]),
        p("12. Integraciones", s["Heading1"]),
        table([
            ["Integracion", "Sprint", "Notas"],
            ["Proveedor SPEI/CLABE virtual", "Sprint 2", "Requiere contrato, sandbox, webhooks y conciliacion."],
            ["Terminal/link tarjeta", "Sprint 2", "Openpay/Conekta/Mercado Pago/banco; confirmar comisiones y flujo."],
            ["Contpaqi", "Futuro", "Exportacion o interfaz, no sustitucion fiscal."],
            ["WhatsApp/OCR", "Futuro", "Solo si se conserva flujo de comprobantes por WhatsApp."],
            ["Camaras/reconocimiento", "Futuro", "Vision posterior, no Sprint 1 ni Sprint 2 inicial."],
            ["Kiosko/ATM", "Futuro", "Depende de hardware y protocolo de integracion."],
        ], [1.7, 1.0, 4.1]),
        p("13. Despliegue requerido", s["Heading1"]),
        bullet_list([
            "Frontend Vite compilado y servido por Nginx o plataforma compatible.",
            "Backend Django con servidor WSGI/ASGI, variables de entorno y HTTPS.",
            "PostgreSQL administrado para produccion; SQLite solo demo/local.",
            "Almacenamiento tipo S3 para fotos, responsivas, comprobantes e identificaciones.",
            "Backups automatizados, logs, monitoreo y proceso de restauracion probado.",
            "Ambientes separados: local, staging y produccion.",
        ], s["BodyText"]),
        p("14. Criterios de aceptacion generales", s["Heading1"]),
        bullet_list([
            "Un contador puede generar Excel y explicar ingresos, egresos, utilidad y saldos.",
            "Un cajero procesa pagos sin ver dashboard ejecutivo.",
            "Un coach puede ver su equipo, formacion, banca y pasar lista.",
            "Un representante confirma efectivo y consulta adeudos de sus alumnos.",
            "Un admin puede mostrar fuga potencial cruzando asistencia y cobranza.",
            "La demo funciona con datos suficientes y pruebas automatizadas pasan.",
        ], s["BodyText"]),
        p("15. Riesgos y mitigaciones", s["Heading1"]),
        table([
            ["Riesgo", "Impacto", "Mitigacion"],
            ["Datos sucios de Excel", "Retrasa migracion y reportes.", "Plantilla, validaciones y piloto por sede."],
            ["Efectivo", "Fugas persistentes.", "Corte diario, responsable, aceptacion y auditoria."],
            ["Reglas ambiguas", "Sistema correcto tecnicamente pero incorrecto operativamente.", "Cerrar reglas de mora, descuentos y topes antes de automatizar."],
            ["Menores de edad", "Riesgo legal y reputacional.", "Permisos, minimo acceso, almacenamiento seguro y consentimiento."],
            ["Pagos reales", "Dependencia externa.", "Sandbox, contrato y pruebas antes de prometer fechas."],
        ], [1.7, 2.1, 3.0]),
    ]
    return story


def governance_story():
    s = styles()
    gantt = ASSETS / "gantt_sprint_plan.png"
    story = [
        p("Documento de gobernanza y ejecucion", s["Title"]),
        p("Futsi Mini ERP - Sprint 1, Sprint 2 y preparacion para despliegue", s["BodyText"]),
        p("Version: 26 de mayo de 2026.", s["Small"]),
        Spacer(1, 0.12 * inch),
        p("Resumen ejecutivo", s["Heading1"]),
        p("La gobernanza del proyecto debe evitar que el MVP crezca hacia un ERP completo antes de validar su valor. La regla central es: primero control y trazabilidad; despues automatizacion, integraciones bancarias, analitica deportiva, camaras y kiosko.", s["Callout"]),
        p("1. Principios de gobierno", s["Heading1"]),
        bullet_list([
            "Priorizar el 80/20: asistencia + cobranza + gastos + descuentos + auditoria.",
            "Mantener permisos por rol y alcance por sede/grupo/representante.",
            "No automatizar reglas ambiguas sin aprobacion operativa.",
            "No manejar datos sensibles sin politica clara de acceso y almacenamiento.",
            "Toda integracion de pagos debe probarse primero en sandbox.",
            "Cada sprint debe cerrar con demo, datos de prueba, checklist y riesgos actualizados.",
        ], s["BodyText"]),
        p("2. Organos y responsabilidades", s["Heading1"]),
        table([
            ["Actor", "Responsabilidad"],
            ["Direccion/dueno", "Prioriza riesgos, aprueba reglas de negocio, valida dashboard ejecutivo."],
            ["Contador", "Define reportes, estado de resultados, exportaciones y conciliacion."],
            ["Coordinador", "Valida flujo de cancha, asistencia, autorizaciones y cierres de sede."],
            ["Cajero", "Valida cobro en ventanilla, efectivo, tarjeta, transferencia y ticket."],
            ["Coach", "Valida pase de lista, equipo, formacion y horas."],
            ["Desarrollo", "Implementa, prueba, documenta, despliega y comunica riesgos."],
        ], [1.6, 5.2]),
        p("3. RACI por modulo", s["Heading1"]),
        table([
            ["Modulo", "Responsable", "Aprueba", "Consulta", "Informado"],
            ["Roles/permisos", "Desarrollo", "Direccion", "Contador/coordinador", "Usuarios"],
            ["Cobranza", "Desarrollo/cajero", "Contador", "Direccion", "Coordinadores"],
            ["Gastos", "Desarrollo/contador", "Direccion", "Coordinador", "Cajero"],
            ["Asistencia", "Desarrollo/coordinador", "Direccion", "Coach", "Contador"],
            ["Reportes", "Desarrollo/contador", "Direccion", "Admin", "Coordinadores"],
            ["Despliegue", "Desarrollo", "Direccion", "Contador", "Usuarios"],
        ], [1.5, 1.5, 1.3, 1.3, 1.2]),
        p("4. Checklist de Sprint 1", s["Heading1"]),
        table([
            ["Elemento", "Estado", "Evidencia / comentario"],
            ["Backend Django/DRF", "Hecho", "API REST con modelos principales y pruebas."],
            ["Frontend React/Tailwind", "Hecho", "Interfaz minimalista por roles."],
            ["Roles", "Hecho", "Admin, direccion, contador, coordinador, cajero, coach, representante."],
            ["Alumnos", "Hecho", "Datos de control, medicos, responsiva, foto, uniforme y pausas."],
            ["Representantes", "Hecho", "Contacto, usuario y CLABE virtual simulada."],
            ["Asistencia", "Hecho", "Pase de lista y deuda al capturar."],
            ["Cobranza", "Hecho demo", "Pagos simulados y estados."],
            ["Gastos", "Hecho base", "Captura/aprobacion basica."],
            ["Portal contador", "Hecho", "Grafica X/Y y Excel."],
            ["Portal coach", "Hecho", "Formacion 4-3-3, banca, asistencia y horas."],
            ["Auditoria", "Parcial", "Modelo existe; falta automatizar captura general."],
            ["Despliegue", "Pendiente", "Debe planearse para Sprint 2."],
        ], [2.0, 1.0, 3.8]),
        PageBreak(),
        p("5. Checklist de Sprint 2", s["Heading1"]),
        table([
            ["Elemento", "Objetivo", "Dependencia"],
            ["Despliegue staging", "Ambiente accesible para pruebas reales.", "Dominio, hosting, variables, BD."],
            ["Despliegue produccion", "Primer piloto operativo.", "Staging validado y backups."],
            ["PostgreSQL", "Base robusta para datos reales.", "Provisionar servicio administrado."],
            ["Archivos S3", "Guardar fotos, responsivas y comprobantes.", "Proveedor cloud y politicas."],
            ["Importacion Excel/Gymforce", "Cargar alumnos, representantes, equipos y pagos historicos.", "Archivos reales limpios."],
            ["Mora y topes", "Automatizar 10 dias + 5% y avisos.", "Reglas aprobadas por direccion."],
            ["Conciliacion SPEI", "Confirmar transferencias por webhook.", "Proveedor financiero."],
            ["Terminal/link real", "Confirmar pagos con tarjeta.", "Proveedor y sandbox."],
            ["Auditoria automatica", "Registrar cambios sensibles.", "Lista de entidades criticas."],
            ["Cierre diario", "Corte por sede y caja.", "Responsables y politica de efectivo."],
        ], [1.8, 3.0, 2.0]),
        p("6. Gantt propuesto", s["Heading1"]),
        Image(str(gantt), width=7.0 * inch, height=3.62 * inch),
        p("7. Politica de cambios", s["Heading1"]),
        numbered_list([
            "Toda solicitud nueva se clasifica como bug, ajuste de Sprint 1, Sprint 2 o futuro.",
            "Cambios que afecten dinero, permisos o datos personales requieren aprobacion de direccion.",
            "Cambios visuales o de copy pueden aprobarse por responsable operativo.",
            "Cambios de reglas de cobro deben documentarse antes de implementarse.",
            "Cada cambio se valida con prueba manual y, si toca logica critica, prueba automatizada.",
        ], s["BodyText"]),
        p("8. Gobierno de datos", s["Heading1"]),
        table([
            ["Dato", "Sensibilidad", "Regla"],
            ["Fotos de menores", "Alta", "Acceso limitado a roles operativos necesarios; almacenar en servicio seguro."],
            ["Responsivas", "Alta", "No descargar masivamente salvo admin/direccion autorizado."],
            ["Datos medicos", "Alta", "Visible solo para quien lo necesita en cancha y administracion."],
            ["Pagos", "Alta", "Visible para contador, direccion, cajero de sede y representante propio."],
            ["Gastos", "Media/alta", "Visible para contador, direccion y aprobadores."],
            ["Auditoria", "Alta", "Solo admin/direccion/contador segun necesidad."],
        ], [1.7, 1.2, 3.9]),
        p("9. Gobierno de seguridad", s["Heading1"]),
        bullet_list([
            "HTTPS obligatorio en produccion.",
            "Tokens y secretos fuera del codigo; variables de entorno.",
            "CORS limitado al dominio del frontend.",
            "Politicas de password y desactivacion de usuarios compartidos.",
            "Backups cifrados y prueba de restauracion.",
            "Logs de acceso y errores sin exponer datos sensibles.",
        ], s["BodyText"]),
        p("10. Plan de despliegue", s["Heading1"]),
        table([
            ["Capa", "Decision recomendada", "Tareas"],
            ["Frontend", "Vite build servido por Nginx o plataforma estatica.", "Variables API, build, cache, dominio."],
            ["Backend", "Django + Gunicorn/Uvicorn detras de Nginx.", "Settings prod, static/media, logs."],
            ["Base", "PostgreSQL administrado.", "Migraciones, backups, usuarios, monitoreo."],
            ["Archivos", "S3 compatible.", "Buckets privados, URLs firmadas, politicas de retencion."],
            ["Pagos", "Proveedor con webhooks.", "Sandbox, firma webhooks, idempotencia, conciliacion."],
            ["Observabilidad", "Logs + errores + metricas.", "Alertas de caida, errores 5xx y webhook fallido."],
        ], [1.1, 2.7, 3.0]),
        p("11. Ceremonias de trabajo", s["Heading1"]),
        table([
            ["Ceremonia", "Frecuencia", "Objetivo"],
            ["Kickoff Sprint", "Inicio", "Confirmar alcance, riesgos y datos necesarios."],
            ["Check diario", "Diario 15 min", "Desbloquear decisiones y validar avance."],
            ["Demo intermedia", "Mitad de sprint", "Validar flujo con usuarios reales."],
            ["Cierre Sprint", "Final", "Demo, checklist, pruebas, riesgos y proximos pasos."],
            ["Revision de datos", "Antes de migrar", "Validar calidad de Excel/Gymforce."],
        ], [1.5, 1.4, 3.9]),
        p("12. Criterios para pasar a produccion piloto", s["Heading1"]),
        bullet_list([
            "Usuarios reales definidos y sin cuentas compartidas.",
            "Reglas de mora, descuento, pago parcial y tope para jugar aprobadas.",
            "Backups y restauracion probados.",
            "Datos migrados y validados contra muestra de Excel.",
            "Pagos reales probados en sandbox y en transaccion pequena.",
            "Cajero, coordinador, contador y coach capacitados.",
            "Plan de contingencia si falla internet o proveedor de pagos.",
        ], s["BodyText"]),
        p("13. Riesgos de gobernanza", s["Heading1"]),
        table([
            ["Riesgo", "Senal temprana", "Accion"],
            ["Scope creep", "Se piden camaras/kiosko antes de cerrar cobranza.", "Regresar a roadmap y clasificar como futuro."],
            ["Reglas cambiantes", "Cada sede opera precios/descuentos distinto.", "Catalogo configurable y aprobacion formal."],
            ["Baja adopcion", "Usuarios vuelven a Excel/WhatsApp.", "Simplificar flujo y capacitar por rol."],
            ["Datos sensibles expuestos", "Roles ven mas informacion de la necesaria.", "Revisar permisos y ocultar campos por rol."],
            ["Conciliacion incompleta", "Pagos quedan en proceso sin seguimiento.", "Alertas y expiracion automatica."],
        ], [1.7, 2.5, 2.6]),
        PageBreak(),
        p("14. Politica de calidad y pruebas", s["Heading1"]),
        p("La calidad no debe medirse solo por que la pantalla cargue. Para este proyecto, la calidad se mide por consistencia operativa: que los datos capturados permitan explicar ingresos, egresos, asistencia y decisiones de autorizacion.", s["BodyText"]),
        table([
            ["Tipo de prueba", "Alcance", "Criterio de salida"],
            ["Unitarias/API", "Permisos, pagos, descuentos, gastos, asistencia y roles.", "Suite automatizada en verde."],
            ["Flujo operativo", "Cajero cobra, representante confirma, contador revisa.", "Flujo demostrado con datos demo y caso real."],
            ["Datos", "Importacion o captura de alumnos/equipos.", "Muestra validada contra archivo origen."],
            ["Seguridad", "Roles y acceso por sede/grupo/representante.", "Usuario no autorizado no ve ni edita datos fuera de alcance."],
            ["Pagos", "Webhook, idempotencia, expiracion y conciliacion.", "Sandbox aprobado antes de produccion."],
            ["Regresion", "Pantallas principales despues de cambios.", "Checklist manual firmado por desarrollo/operacion."],
        ], [1.45, 3.2, 2.15]),
        p("15. Plan de capacitacion", s["Heading1"]),
        table([
            ["Perfil", "Capacitacion minima", "Material requerido"],
            ["Cajero", "Crear solicitud de pago, elegir metodo, validar estado, explicar efectivo/transferencia/link.", "Guia de cobro y casos de error."],
            ["Coordinador", "Crear sesion, pasar lista, interpretar adeudos, cerrar asistencia y justificar excepciones.", "Guia de cancha y reglas de autorizacion."],
            ["Coach", "Ver equipo, pasar lista, revisar alertas medicas y registrar horas.", "Guia de coach por grupo."],
            ["Contador", "Exportar Excel, leer dashboard, revisar gastos y conciliar pagos.", "Guia contable y diccionario de campos."],
            ["Direccion", "Interpretar tablero de fugas, utilidad, sedes y riesgos.", "Resumen ejecutivo y checklist semanal."],
            ["Representante", "Consultar CLABE, adeudos, link, aceptar efectivo.", "Guia corta desde portal cliente."],
        ], [1.3, 3.5, 2.0]),
        p("16. Metricas de gobierno", s["Heading1"]),
        table([
            ["Metrica", "Por que importa", "Frecuencia"],
            ["Alumnos/equipos con asistencia y adeudo", "Detecta fuga directa entre operacion y cobranza.", "Diario/semanal"],
            ["Pagos en efectivo pendientes de aceptacion", "Reduce disputa entre cajero y cliente.", "Diario"],
            ["Transferencias en proceso mayor a 72 horas", "Evita que pagos simulados queden sin seguimiento.", "Diario"],
            ["Descuentos por coordinador/sede", "Detecta uso irregular de descuentos.", "Semanal"],
            ["Gastos pendientes o rechazados", "Controla egresos no aprobados.", "Semanal"],
            ["Utilidad por sede", "Permite comparar operacion real.", "Semanal/mensual"],
            ["Horas de coach vs alumnos/asistencia", "Controla nomina deportiva.", "Semanal"],
            ["Tickets reabiertos o cambios post-cierre", "Mide salud de trazabilidad.", "Semanal"],
        ], [1.8, 3.3, 1.7]),
        p("17. Gobierno de despliegue", s["Heading1"]),
        numbered_list([
            "Preparar ambiente staging con datos anonimizados o demo realista.",
            "Ejecutar migraciones sobre PostgreSQL en staging.",
            "Configurar almacenamiento de archivos y permisos privados.",
            "Configurar dominio, HTTPS, CORS y variables de entorno.",
            "Ejecutar pruebas de smoke: login, dashboard, caja, cliente, coach, contador.",
            "Realizar respaldo inicial antes de importar datos reales.",
            "Hacer piloto con una sede y una semana operativa.",
            "Revisar diferencias contra Excel antes de retirar cualquier flujo manual.",
        ], s["BodyText"]),
        p("18. Politica de incidentes", s["Heading1"]),
        table([
            ["Incidente", "Severidad", "Respuesta"],
            ["No se puede cobrar", "Alta", "Usar contingencia manual, registrar folio y cargar despues con evidencia."],
            ["Webhook de pago falla", "Alta", "Reintentar, revisar logs, marcar pago en revision sin darlo por conciliado."],
            ["Datos sensibles expuestos", "Critica", "Bloquear acceso, revisar logs, corregir permisos y notificar responsables."],
            ["Error de asistencia", "Media", "Solicitar reapertura con aprobacion y dejar bitacora."],
            ["Gasto duplicado", "Media", "Marcar en revision, validar comprobante y aprobar/rechazar."],
            ["Caida de sistema", "Alta", "Activar procedimiento de contingencia y revisar infraestructura."],
        ], [1.8, 1.1, 3.9]),
        p("19. Preguntas abiertas para cerrar antes de produccion", s["Heading1"]),
        bullet_list([
            "Proveedor de pagos elegido: banco, STP, Openpay, Conekta, Mercado Pago u otro.",
            "Formato final del estado de resultados que recibe el contador.",
            "Politica exacta de mora, topes y excepciones para academia y torneos.",
            "Reglas de reembolso, congelamiento de saldo y transferencia entre hermanos.",
            "Alcance de inventario de uniformes: solo venta o stock por sede/talla.",
            "Politica de conservacion de datos de menores, documentos e identificaciones.",
            "Proceso de baja de usuarios y prohibicion de cuentas compartidas.",
            "Responsable operativo de aprobar reaperturas de cierres.",
        ], s["BodyText"]),
        p("20. Definicion de listo y definicion de terminado", s["Heading1"]),
        table([
            ["Concepto", "Definicion"],
            ["Listo para desarrollo", "Regla entendida, datos requeridos disponibles, responsable de aprobacion identificado y criterio de aceptacion escrito."],
            ["Terminado para demo", "Funciona localmente, tiene datos demo, no rompe permisos y puede explicarse en junta."],
            ["Terminado para produccion", "Probado en staging, con backup, monitoreo, permisos revisados, capacitacion y plan de contingencia."],
            ["Terminado para retirar Excel", "Una semana operativa conciliada contra Excel/Gymforce sin diferencias materiales."],
        ], [2.1, 4.7]),
    ]
    return story


def build_docx_from_sections(path: Path, title: str, subtitle: str, sections: list[tuple[str, list[str]]]):
    doc = Document()
    docx_style(doc)
    docx_title(doc, title, subtitle)
    for heading, items in sections:
        doc.add_heading(heading, level=1)
        for item in items:
            if item.startswith("- "):
                doc.add_paragraph(item[2:], style="List Bullet")
            else:
                doc.add_paragraph(item)
    doc.save(path)


def build_editable_docx():
    build_docx_from_sections(
        DOCS / "Futsi_Especificacion_Requerimientos.docx",
        "Especificacion de requerimientos de software",
        "Futsi Mini ERP - documento formal ampliado",
        [
            ("Resumen ejecutivo", ["El sistema crea una fuente confiable de verdad para cruzar asistencia, cobranza, descuentos, gastos y trazabilidad por sede. No sustituye Contpaqi ni es un ERP fiscal completo."]),
            ("Alcance Sprint 1", ["- Login por roles.", "- Sedes, alumnos, representantes y equipos base.", "- Asistencia, cobranza, gastos, descuentos y dashboards.", "- Portal contador y portal coach.", "- Documentacion, pruebas y datos demo."]),
            ("Requerimientos principales", ["- Controlar cargos, pagos parciales y saldos.", "- Simular CLABE virtual, efectivo aceptado, terminal y link.", "- Exportar Excel contable.", "- Mostrar grafica financiera con eje X/Y.", "- Gestionar datos de alumno: foto, responsiva, medico, uniforme, pausas.", "- Registrar horas de coach y formacion del equipo."]),
            ("Sprint 2 recomendado", ["- Despliegue staging/produccion.", "- PostgreSQL y almacenamiento S3.", "- Importacion Excel/Gymforce.", "- SPEI/terminal reales.", "- Mora, topes y avisos.", "- Auditoria automatica y cierre diario."]),
        ],
    )
    build_docx_from_sections(
        DOCS / "Futsi_Gobernanza_Roadmap.docx",
        "Gobernanza, roadmap y ejecucion",
        "Futsi Mini ERP - documento formal ampliado",
        [
            ("Principio", ["Primero control y trazabilidad; despues automatizacion, integraciones y analitica avanzada. El MVP debe demostrar que asistencia, cobranza, descuentos, gastos y auditoria reducen fugas sin construir un ERP completo."]),
            ("Gobierno", ["- Direccion aprueba reglas y prioridades.", "- Contador valida reportes y conciliacion.", "- Coordinador valida cancha y cierres.", "- Cajero valida cobros.", "- Coach valida asistencia y horas.", "- Desarrollo implementa, prueba y despliega.", "- Cambios que afecten dinero, permisos o datos personales requieren aprobacion formal."]),
            ("Checklist Sprint 1", ["- Backend Django/DRF implementado.", "- Frontend React/Tailwind implementado.", "- Roles implementados.", "- Alumnos y representantes implementados.", "- Asistencia implementada.", "- Cobranza simulada implementada.", "- Gastos base implementados.", "- Portal contador implementado.", "- Portal coach implementado.", "- Auditoria automatica parcial.", "- Despliegue pendiente."]),
            ("Sprint 2", ["- Despliegue staging y produccion.", "- PostgreSQL.", "- S3 compatible.", "- Importacion Excel/Gymforce.", "- Mora y topes.", "- SPEI/terminal reales.", "- Auditoria automatica.", "- Cierre diario.", "- Piloto por sede."]),
            ("Despliegue", ["- Frontend Vite.", "- Backend Django.", "- PostgreSQL.", "- S3 compatible.", "- HTTPS.", "- Backups.", "- Monitoreo.", "- Sandbox de pagos.", "- Variables de entorno.", "- CORS cerrado.", "- Logs y alertas."]),
            ("Calidad", ["- Pruebas API.", "- Pruebas manuales por rol.", "- Validacion de datos.", "- Revision de seguridad.", "- Smoke test antes de liberar.", "- Checklist firmado por responsable."]),
            ("Capacitacion", ["- Cajero: pagos y estados.", "- Coordinador: asistencia y excepciones.", "- Coach: lista, alertas y horas.", "- Contador: Excel y dashboard.", "- Direccion: tablero y riesgos.", "- Representante: adeudos y confirmaciones."]),
            ("Riesgos", ["- Scope creep.", "- Datos sucios.", "- Efectivo.", "- Reglas ambiguas.", "- Datos sensibles.", "- Dependencia de proveedor de pagos.", "- Baja adopcion."]),
        ],
    )


def main():
    pdf(DOCS / "Futsi_Especificacion_Requerimientos.pdf", srs_story())
    pdf(DOCS / "Futsi_Gobernanza_Roadmap.pdf", governance_story())
    build_editable_docx()
    print("Formal documents rebuilt")


if __name__ == "__main__":
    main()
