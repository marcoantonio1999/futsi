from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    Image,
    ListFlowable,
    ListItem,
    PageBreak,
    Paragraph,
    SimpleDocTemplate,
    Spacer,
    Table,
    TableStyle,
)
from pypdf import PdfReader, PdfWriter


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
ASSETS = DOCS / "assets"
DOCS.mkdir(exist_ok=True)

ACCENT = colors.HexColor("#14532d")
BORDER = colors.HexColor("#cbd5e1")
FILL = colors.HexColor("#dcfce7")
INK = colors.HexColor("#0f172a")
MUTED = colors.HexColor("#52525b")


def make_styles():
    styles = getSampleStyleSheet()
    styles["Title"].fontName = "Helvetica-Bold"
    styles["Title"].fontSize = 21
    styles["Title"].leading = 25
    styles["Title"].textColor = INK
    styles["Heading1"].fontName = "Helvetica-Bold"
    styles["Heading1"].fontSize = 15
    styles["Heading1"].leading = 18
    styles["Heading1"].spaceBefore = 14
    styles["Heading1"].spaceAfter = 7
    styles["Heading1"].textColor = ACCENT
    styles["Heading2"].fontName = "Helvetica-Bold"
    styles["Heading2"].fontSize = 12
    styles["Heading2"].leading = 15
    styles["Heading2"].spaceBefore = 10
    styles["Heading2"].spaceAfter = 5
    styles["Heading2"].textColor = INK
    styles["BodyText"].fontName = "Helvetica"
    styles["BodyText"].fontSize = 9.2
    styles["BodyText"].leading = 12
    styles["BodyText"].spaceAfter = 5
    styles.add(ParagraphStyle(name="Small", parent=styles["BodyText"], fontSize=8, leading=10, textColor=MUTED))
    styles.add(ParagraphStyle(name="Callout", parent=styles["BodyText"], backColor=colors.HexColor("#f0fdf4"), borderColor=colors.HexColor("#bbf7d0"), borderWidth=0.6, borderPadding=7, spaceBefore=6, spaceAfter=8))
    return styles


def P(text: str, style):
    return Paragraph(text, style)


def bullets(items: list[str], style):
    return ListFlowable([ListItem(P(item, style)) for item in items], bulletType="bullet", leftIndent=16)


def nums(items: list[str], style):
    return ListFlowable([ListItem(P(item, style)) for item in items], bulletType="1", leftIndent=18)


def tbl(data: list[list[str]], widths: list[float]):
    s = make_styles()
    table = Table([[P(str(cell), s["Small"]) for cell in row] for row in data], colWidths=[w * inch for w in widths], repeatRows=1)
    table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), FILL),
        ("GRID", (0, 0), (-1, -1), 0.35, BORDER),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
    ]))
    return table


def build_pdf(path: Path, story: list):
    def footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(MUTED)
        canvas.drawString(0.65 * inch, 0.38 * inch, "Futsi Mini ERP - documentacion Sprint 2")
        canvas.drawRightString(7.85 * inch, 0.38 * inch, f"Pagina {doc.page}")
        canvas.restoreState()

    doc = SimpleDocTemplate(
        str(path),
        pagesize=letter,
        rightMargin=0.65 * inch,
        leftMargin=0.65 * inch,
        topMargin=0.62 * inch,
        bottomMargin=0.62 * inch,
    )
    doc.build(story, onFirstPage=footer, onLaterPages=footer)


def build_docx(path: Path, title: str, sections: list[tuple[str, list[str]]]):
    doc = Document()
    section = doc.sections[0]
    section.top_margin = section.bottom_margin = section.left_margin = section.right_margin = Inches(1)
    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    title_paragraph = doc.add_paragraph()
    run = title_paragraph.add_run(title)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.bold = True
    run.font.color.rgb = RGBColor.from_string("0F172A")
    for heading, items in sections:
        doc.add_heading(heading, level=1)
        for item in items:
            if item.startswith("- "):
                doc.add_paragraph(item[2:], style="List Bullet")
            else:
                doc.add_paragraph(item)
    doc.save(path)


def code_documentation_story():
    s = make_styles()
    return [
        P("Documentacion tecnica del codigo", s["Title"]),
        P("Futsi Mini ERP - React/Vite/Tailwind + Django REST Framework", s["BodyText"]),
        P("Actualizada al 27 de mayo de 2026.", s["Small"]),
        Spacer(1, 0.12 * inch),
        P("1. Resumen tecnico", s["Heading1"]),
        P("El proyecto esta dividido en dos aplicaciones: backend Django en <b>back/</b> y frontend React en <b>front/</b>. La API expone recursos REST con Django REST Framework y autenticacion por Token. El frontend consume la API con fetch, renderiza portales por rol y usa Tailwind para la interfaz minimalista.", s["BodyText"]),
        P("Estructura principal", s["Heading1"]),
        tbl([
            ["Ruta", "Proposito"],
            ["back/futsi_api/settings.py", "Configuracion Django: apps, DB SQLite/PostgreSQL, CORS, REST Framework, timezone, media/static."],
            ["back/futsi_api/urls.py", "Router REST y endpoints auth/login/logout/me."],
            ["back/core/models.py", "Modelo de dominio: usuarios, sedes, alumnos, asistencia, cargos, pagos, descuentos, gastos, coach logs, cierres y auditoria."],
            ["back/core/serializers.py", "Serializacion DRF, calculo de saldos, estados de cargos y simulacion de pagos."],
            ["back/core/views.py", "ViewSets, filtros por rol, acciones de cierre, confirmacion de efectivo, webhook simulado y aprobaciones."],
            ["back/core/permissions.py", "Permisos por rol: admin, operaciones, cajero, coach y representante."],
            ["back/core/management/commands/seed_demo.py", "Carga reproducible de datos demo y usuarios de prueba."],
            ["front/src/main.tsx", "Aplicacion React completa: login, portales, dashboards, formularios, mapas, graficas y export Excel."],
            ["front/src/styles.css", "Entrada Tailwind y estilos globales."],
            ["tools/*.py", "Generadores de documentacion, ERD, Gantt y PDFs."],
        ], [2.0, 4.8]),
        P("2. Stack y dependencias", s["Heading1"]),
        tbl([
            ["Capa", "Tecnologia", "Notas"],
            ["Frontend", "React 19, Vite 7, TypeScript, Tailwind 4", "SPA con portales por rol y componentes locales."],
            ["Mapas", "Leaflet + OpenStreetMap", "Mapa real de sedes con coordenadas."],
            ["Iconos", "lucide-react", "Botones y acciones."],
            ["Backend", "Django 5.2, DRF, django-cors-headers", "API REST con auth token."],
            ["Base local", "SQLite", "Solo desarrollo/demo."],
            ["Base produccion", "PostgreSQL", "Ya soportado por variables en settings.py."],
            ["Docs", "reportlab, python-docx, Pillow", "Generacion de PDFs, DOCX e imagenes."],
        ], [1.3, 2.1, 3.4]),
        P("3. Configuracion backend", s["Heading1"]),
        tbl([
            ["Variable", "Uso", "Ejemplo"],
            ["DJANGO_SECRET_KEY", "Clave secreta Django.", "valor largo generado"],
            ["DJANGO_DEBUG", "Activar/desactivar debug.", "false en produccion"],
            ["DJANGO_ALLOWED_HOSTS", "Hosts permitidos.", "api.midominio.com"],
            ["CORS_ALLOWED_ORIGINS", "Origenes frontend permitidos.", "https://app.midominio.com"],
            ["DB_ENGINE", "Selecciona sqlite o postgres.", "postgres"],
            ["POSTGRES_DB", "Nombre de base.", "futsi"],
            ["POSTGRES_USER", "Usuario BD.", "futsi_app"],
            ["POSTGRES_PASSWORD", "Password BD.", "secreto"],
            ["POSTGRES_HOST", "Host BD.", "db.internal"],
            ["POSTGRES_PORT", "Puerto BD.", "5432"],
        ], [1.7, 2.6, 2.5]),
        P("4. Modelo de dominio", s["Heading1"]),
        bullets([
            "User extiende AbstractUser y agrega role, primary_site, phone, avatar_url, coach_group_name y coach_hourly_rate.",
            "Site representa sedes con coordenadas, estado activo y ventana de edicion.",
            "Guardian representa padre/tutor/representante y genera virtual_clabe si falta.",
            "Student guarda datos operativos y sensibles: foto, responsiva, medico, uniforme y pausas.",
            "Tournament, Team, Player y Round preparan torneos adultos y jornadas.",
            "AttendanceSession y AttendanceRecord capturan asistencia y si existia adeudo.",
            "Charge, Payment y Discount modelan cobranza, pagos parciales, descuentos y conciliacion simulada.",
            "Expense modela gastos por sede y aprobacion.",
            "CoachWorkLog captura horas de coach con tarifa snapshot.",
            "DailyClosure y AuditLog preparan cierres y auditoria.",
        ], s["BodyText"]),
        P("5. API REST", s["Heading1"]),
        tbl([
            ["Endpoint", "Recurso", "Notas de permiso"],
            ["/api/auth/login/", "Login", "Publico, devuelve token y usuario."],
            ["/api/auth/me/", "Perfil actual", "GET/PATCH para contacto y perfil representante."],
            ["/api/users/", "Usuarios", "Solo admin/owner."],
            ["/api/sites/", "Sedes", "Lectura autenticada; escritura admin/owner."],
            ["/api/students/", "Alumnos", "Filtra guardian, cajero y coach por alcance."],
            ["/api/attendance-sessions/", "Sesiones", "Operaciones y coach."],
            ["/api/attendance-records/", "Asistencia", "Operaciones, coach y lectura guardian."],
            ["/api/charges/", "Cargos", "Operaciones; lectura cajero/guardian acotada."],
            ["/api/payments/", "Pagos", "Operaciones, cajero y guardian para acciones permitidas."],
            ["/api/payments/{id}/confirm-cash/", "Confirmar efectivo", "Representante propio u operaciones."],
            ["/api/payments/{id}/simulate-webhook/", "Confirmar SPEI/link demo", "Simulacion de proveedor."],
            ["/api/discounts/", "Descuentos", "Operaciones; guardian lectura acotada."],
            ["/api/expenses/", "Gastos", "Operaciones."],
            ["/api/coach-work-logs/", "Horas coach", "Operaciones y coach; coach solo sus registros."],
        ], [2.2, 1.8, 2.8]),
        PageBreak(),
        P("6. Logica de cobranza", s["Heading1"]),
        bullets([
            "charge_balance calcula monto pendiente restando pagos registrados/conciliados y descuentos aprobados.",
            "sync_charge_status marca paid, partial o pending segun saldo.",
            "PaymentSerializer asigna site/student/team desde el cargo para evitar capturas inconsistentes.",
            "Transferencia queda en processing, referencia CLABE y expiracion 72 horas.",
            "Efectivo queda awaiting_confirmation hasta aceptacion del representante.",
            "Tarjeta terminal queda registered automaticamente en demo.",
            "Link de tarjeta queda processing y puede simularse con webhook.",
            "Los pagos en processing no reducen saldo hasta confirmarse.",
        ], s["BodyText"]),
        P("7. Frontend", s["Heading1"]),
        tbl([
            ["Componente / funcion", "Responsabilidad"],
            ["App", "Carga usuario, decide portal segun rol y administra token/data."],
            ["LoginScreen", "Formulario de acceso."],
            ["DashboardPanel", "Vista admin/direccion con metricas, mapa, graficas y alertas."],
            ["AccountingPortal", "Vista contador con grafica X/Y y export Excel."],
            ["CashierPortal", "Vista cajero para procesar pagos por metodo."],
            ["GuardianPortal", "Vista representante: adeudos, CLABE, links y efectivo."],
            ["CoachPortal", "Vista coach: formacion, banca, asistencia y horas."],
            ["AttendancePanel", "Pase de lista operativo por sede/grupo."],
            ["BillingPanel", "Cargos, pagos, descuentos y aprobaciones."],
            ["ExpensesPanel", "Captura y aprobacion de gastos."],
            ["StudentsPanel", "Alta/edicion/filtros de alumnos y datos sensibles."],
            ["exportAccountingWorkbook", "Genera Excel XML descargable desde navegador."],
        ], [2.2, 4.6]),
        P("8. Comandos de desarrollo", s["Heading1"]),
        tbl([
            ["Accion", "Comando"],
            ["Instalar backend", "cd back; python -m venv .venv; .venv\\Scripts\\pip install -r requirements.txt"],
            ["Migrar backend", "cd back; .venv\\Scripts\\python manage.py migrate"],
            ["Seed demo", "cd back; .venv\\Scripts\\python manage.py seed_demo --reset"],
            ["Correr API", "cd back; .venv\\Scripts\\python manage.py runserver 127.0.0.1:8000"],
            ["Instalar frontend", "cd front; npm install"],
            ["Correr frontend", "cd front; npm run dev -- --host 127.0.0.1"],
            ["Build frontend", "cd front; npm run build"],
            ["Tests backend", "cd back; .venv\\Scripts\\python manage.py test"],
        ], [1.7, 5.1]),
        P("9. Datos demo y credenciales", s["Heading1"]),
        tbl([
            ["Rol", "Usuario", "Password"],
            ["Admin", "admin", "admin12345"],
            ["Contador", "contador", "demo12345"],
            ["Coordinador Roma", "coordinador.roma", "demo12345"],
            ["Cajero Roma", "caja.roma", "demo12345"],
            ["Cajero Coyoacan", "caja.coyoacan", "demo12345"],
            ["Coach Roma", "coach.roma", "demo12345"],
            ["Representante", "padre.laura / padre.roberto / padre.daniela / padre.jorge", "familia12345"],
        ], [1.4, 3.2, 1.4]),
        P("10. Testing actual", s["Heading1"]),
        bullets([
            "Flujo integral: asistencia + pago + descuento + gasto.",
            "Representante solo ve sus alumnos y no puede crear cargos.",
            "Representante puede editar datos de contacto/perfil.",
            "Campos de control del alumno son editables.",
            "Cajero solo ve su sede y puede procesar pagos.",
            "Simulacion de transferencia, efectivo y webhook/link.",
            "Coach solo ve su grupo y puede registrar asistencia/horas sin crear gastos.",
        ], s["BodyText"]),
        P("11. Deuda tecnica y Sprint 2 tecnico", s["Heading1"]),
        tbl([
            ["Tema", "Necesidad"],
            ["Auditoria", "Automatizar AuditLog con signals o capa de servicio para pagos, descuentos, gastos, alumnos y cierres."],
            ["Settings produccion", "Separar settings dev/prod o endurecer por env vars."],
            ["Archivos", "Mover media local a S3 compatible con URLs privadas/firmadas."],
            ["Pagos", "Integrar proveedor real y validar firma/idempotencia de webhooks."],
            ["Mora", "Job programado para generar penalizaciones y avisos."],
            ["Importacion", "Comando o pantalla de carga Excel con validaciones y reporte de errores."],
            ["Permisos", "Endurecer escrituras por rol y agregar pruebas por endpoint."],
            ["Frontend", "Separar componentes grandes de main.tsx a modulos por portal."],
            ["Observabilidad", "Logs estructurados, Sentry o similar y health checks."],
        ], [1.6, 5.2]),
        P("12. Recomendacion de refactor", s["Heading1"]),
        nums([
            "Separar front/src/main.tsx en modules: api.ts, types.ts, portals/, panels/, components/.",
            "Crear servicios backend para pagos, descuentos, auditoria y cierres.",
            "Agregar permisos por objeto donde haya mas datos reales.",
            "Crear OpenAPI/Swagger para documentar endpoints.",
            "Agregar factories/fixtures de prueba para no depender solo de seed_demo.",
        ], s["BodyText"]),
    ]


def business_documentation_story():
    s = make_styles()
    return [
        P("Documentacion de negocio", s["Title"]),
        P("Futsi Mini ERP - procesos, roles, reglas y operacion esperada", s["BodyText"]),
        P("Actualizada al 27 de mayo de 2026.", s["Small"]),
        Spacer(1, 0.12 * inch),
        P("1. Proposito", s["Heading1"]),
        P("Este documento traduce el sistema a lenguaje operativo para direccion, contador, coordinadores, cajeros, coaches y representantes. Su objetivo es explicar como debe operar el negocio con la herramienta, que controles existen y que falta definir antes de produccion.", s["Callout"]),
        P("2. Principio 80/20", s["Heading1"]),
        P("El valor inicial no esta en tener muchas pantallas, sino en saber con confianza: quien asistio o jugo, quien debia pagar, quien pago, quien registro el pago, que descuento se aplico, que gasto se reporto y que informacion quedo cerrada.", s["BodyText"]),
        P("3. Roles de negocio", s["Heading1"]),
        tbl([
            ["Rol", "Responsabilidad", "Indicadores que revisa"],
            ["Direccion", "Detectar fugas, evaluar sedes, decidir reglas y aprobar prioridades.", "Ingresos, egresos, utilidad, adeudos, descuentos, gastos pendientes."],
            ["Contador", "Preparar estado de resultados, revisar pagos/no pagos y conciliacion.", "Excel contable, pagos por metodo, gastos, saldos, asistencia con adeudo."],
            ["Coordinador", "Operar sede y cancha; autorizar excepciones.", "Asistencia, alumnos con adeudo, cierres, incidencias."],
            ["Cajero", "Procesar pagos en ventanilla.", "Cargos abiertos, pagos en proceso, efectivo pendiente de aceptacion."],
            ["Coach", "Pasar lista y registrar horas.", "Equipo, alertas medicas, adeudos visibles, horas registradas."],
            ["Representante", "Pagar y confirmar informacion propia.", "Adeudos, CLABE, links, efectivo pendiente."],
        ], [1.3, 3.1, 2.4]),
        P("4. Proceso de academia infantil", s["Heading1"]),
        nums([
            "Alumno entra en prueba por 2 semanas/2 clases; puede extenderse hasta 4 semanas con autorizacion.",
            "Si continua, se registra como activo, con representante, sede, grupo, foto, responsiva y datos medicos.",
            "Se genera cargo mensual por mes completo y cargos adicionales como uniforme o torneo escolar.",
            "El representante ve adeudos y medios de pago.",
            "El coordinador/coach pasa lista y ve si el alumno tiene adeudo, pausa, lesion o alerta medica.",
            "Si hay adeudo y se permite entrenar, queda registrado como asistencia con adeudo.",
            "Contador revisa alumnos pagados/no pagados y saldos por sede.",
        ], s["BodyText"]),
        P("5. Proceso de torneos de adultos", s["Heading1"]),
        nums([
            "Torneo se configura por sede, tipo de cobro y duracion esperada.",
            "Puede cobrarse torneo completo durante primeras 3 jornadas o por jornada semanal.",
            "Una jornada equivale a una semana de juegos; puede existir doble jornada.",
            "Torneo dura 12 jornadas mas liguilla.",
            "Equipo tiene representante y jugadores; algunos jugadores pueden pagar individualmente.",
            "Si equipo no paga puede jugar con decision del coordinador, pero se genera mora y debe existir tope.",
            "Contador y direccion revisan equipos que jugaron sin pago o con pago parcial.",
        ], s["BodyText"]),
        P("6. Cobranza y medios de pago", s["Heading1"]),
        tbl([
            ["Metodo", "Flujo ideal", "Control"],
            ["Transferencia", "Cliente usa CLABE virtual; cajero registra intencion/monto; sistema confirma por webhook.", "Estado en proceso hasta recibir pago; si no llega en 72h vuelve a adeudo."],
            ["Efectivo", "Cajero registra solicitud; representante acepta desde portal.", "Ambas partes confirman que dinero fue entregado/recibido."],
            ["Tarjeta terminal", "Cajero crea cobro; terminal/proveedor confirma.", "Pago queda registrado automaticamente."],
            ["Link tarjeta", "Cajero o portal envia link; cliente paga desde app.", "Pago en proceso hasta confirmacion del proveedor."],
            ["Cortesia/beca", "Solo con autorizacion.", "Debe quedar motivo y responsable."],
        ], [1.3, 3.4, 2.1]),
        P("7. Reglas de cobranza pendientes de cerrar", s["Heading1"]),
        tbl([
            ["Regla", "Propuesta actual", "Decision pendiente"],
            ["Mora academia", "5% despues de 10 dias.", "Confirmar porcentaje, redondeo y excepciones."],
            ["Avisos", "1 papas, 2 nino, 3 no juega, 4 no entrena.", "Definir si es automatico por sistema."],
            ["Torneos", "Mora por equipo/jornada impaga.", "Definir tope y bloqueo."],
            ["Pagos parciales", "Permitidos.", "Definir aplicacion al adeudo mas antiguo."],
            ["Reembolsos", "Existen.", "Definir flujo, aprobador y asiento operativo."],
            ["Pausas", "Autorizadas por lesion/viaje.", "Definir impacto en mensualidad."],
        ], [1.4, 2.4, 3.0]),
        PageBreak(),
        P("8. Descuentos", s["Heading1"]),
        tbl([
            ["Descuento", "Uso", "Control requerido"],
            ["Hermanos", "Normalmente 15%.", "Aplicar por representante/familia y vigencia."],
            ["Referido", "Promocion comercial.", "Validar condicion y periodo."],
            ["Efectivo", "Si negocio decide mantenerlo.", "Alto riesgo; debe estar autorizado."],
            ["Lesion", "Ajuste por impedimento.", "Evidencia o aprobacion y cruce con asistencia."],
            ["Viaje largo", "Pausa/descuento temporal.", "Fechas de inicio/fin y autorizador."],
            ["Especial", "Caso no estandar.", "Aprobacion admin/direccion y comentario obligatorio."],
        ], [1.4, 2.5, 2.9]),
        P("9. Gastos y nomina", s["Heading1"]),
        bullets([
            "Coordinador solicita o captura gastos operativos segun politica.",
            "Contador revisa soporte y direccion paga/aprueba gastos sensibles.",
            "Coaches ganan diferente y por hora; el portal coach registra horas con tarifa snapshot.",
            "Sprint 2 debe cruzar horas de coach contra asistencia real y grupos atendidos.",
            "Gastos recurrentes deben diferenciarse de gastos extraordinarios.",
            "Todo gasto relevante debe tener sede, categoria, fecha, monto, proveedor/persona, comprobante y estatus.",
        ], s["BodyText"]),
        P("10. Cierres operativos", s["Heading1"]),
        nums([
            "Al final del dia, sede cierra asistencia, pagos, efectivo, gastos e incidencias.",
            "Despues del cierre, la informacion no debe editarse libremente.",
            "Correcciones posteriores requieren solicitud, aprobacion y bitacora.",
            "Contador revisa diferencias contra pagos, banco y efectivo.",
            "Direccion ve alertas de fugas o gastos irregulares.",
        ], s["BodyText"]),
        P("11. Reportes de negocio", s["Heading1"]),
        tbl([
            ["Reporte", "Usuario", "Uso"],
            ["Estado de resultados por sede", "Contador/direccion", "Comparar utilidad e identificar sedes problemáticas."],
            ["Cobranza abierta", "Cajero/contador", "Dar seguimiento a adeudos y pagos parciales."],
            ["Asistencia con adeudo", "Direccion/coordinador", "Detectar ninos/equipos atendidos sin pago."],
            ["Descuentos por coordinador", "Direccion", "Detectar descuentos anormales."],
            ["Gastos pendientes", "Contador/direccion", "Control de aprobaciones y comprobantes."],
            ["Horas coach vs alumnos", "Direccion/contador", "Control de nomina deportiva."],
            ["Torneos extendidos", "Direccion", "Medir si duracion reduce rentabilidad."],
        ], [1.8, 1.8, 3.2]),
        P("12. Sprint 2 de negocio", s["Heading1"]),
        tbl([
            ["Tema", "Resultado esperado"],
            ["Despliegue", "Sistema usable por usuarios reales en staging y piloto."],
            ["Migracion", "Alumnos, representantes, equipos y cargos iniciales cargados desde Excel/Gymforce."],
            ["Reglas", "Mora, avisos, topes, reembolsos, pausas y descuentos aprobados."],
            ["Pagos reales", "CLABE/SPEI, terminal/link y webhooks funcionando con proveedor."],
            ["Cierre diario", "Corte de sede, caja y evidencia de efectivo."],
            ["Auditoria", "Cambios sensibles registrados automaticamente."],
            ["Reportes deportivos", "Posiciones, categorias, resultados y metricas por coach."],
        ], [1.7, 5.1]),
        P("13. Criterios para operar piloto", s["Heading1"]),
        bullets([
            "Una sede seleccionada con coordinador, cajero, coach y contador capacitados.",
            "Datos migrados y validados contra archivo origen.",
            "Reglas de cobro aprobadas por direccion.",
            "Proveedor de pagos probado o flujo manual temporal aceptado.",
            "Plan de contingencia si falla internet, terminal o API.",
            "Cierre diario definido y responsable asignado.",
            "Reporte semanal revisado con contador y direccion.",
        ], s["BodyText"]),
        P("14. Cambios culturales necesarios", s["Heading1"]),
        bullets([
            "Eliminar cuentas compartidas como ocurre en Gymforce.",
            "Dejar de usar Excel como fuente primaria y usarlo solo como exportacion/reporte.",
            "Registrar excepciones en sistema, no solo por WhatsApp.",
            "Cajero, coordinador y coach deben operar con su propio usuario.",
            "Direccion debe revisar tableros semanalmente para que el control tenga efecto.",
        ], s["BodyText"]),
    ]


def governance_update_story():
    s = make_styles()
    gantt = ASSETS / "gantt_sprint_plan.png"
    return [
        P("Gobernanza Sprint 2 - actualizacion", s["Title"]),
        P("Complemento al documento de gobernanza existente. Fecha: 27 de mayo de 2026.", s["BodyText"]),
        P("Objetivo Sprint 2", s["Heading1"]),
        P("Convertir la demo funcional de Sprint 1 en un piloto controlado: desplegable, con datos reales, reglas de negocio cerradas, pagos reales o semirreales, auditoria automatica y reportes que el contador pueda usar semanalmente.", s["Callout"]),
        P("Backlog Sprint 2 priorizado", s["Heading1"]),
        tbl([
            ["Prioridad", "Epica", "Entregable"],
            ["P0", "Despliegue", "Staging + produccion piloto con PostgreSQL, HTTPS, backups y variables."],
            ["P0", "Datos reales", "Importador Excel/Gymforce con reporte de errores y limpieza."],
            ["P0", "Reglas de cobro", "Mora 10 dias/5%, avisos, topes y pagos parciales."],
            ["P1", "Pagos reales", "Proveedor SPEI/CLABE y tarjeta/link con webhooks."],
            ["P1", "Auditoria", "Registro automatico de cambios sensibles."],
            ["P1", "Cierre diario", "Caja, efectivo, pagos, gastos e incidencias por sede."],
            ["P2", "Nomina coaches", "Horas vs grupos/asistencia y aprobacion contable."],
            ["P2", "Deportivo", "Tablas, posiciones, resultados y metricas por categoria/coach."],
        ], [0.8, 1.8, 4.2]),
        P("Gantt de referencia", s["Heading1"]),
        Image(str(gantt), width=7.0 * inch, height=3.62 * inch),
        PageBreak(),
        P("Checklist Sprint 2", s["Heading1"]),
        tbl([
            ["Checklist", "Responsable", "Estado inicial"],
            ["Dominio/subdominios definidos.", "Direccion/desarrollo", "Pendiente"],
            ["Hosting backend seleccionado.", "Desarrollo", "Pendiente"],
            ["PostgreSQL provisionado.", "Desarrollo", "Pendiente"],
            ["Storage de archivos privado.", "Desarrollo", "Pendiente"],
            ["Proveedor de pagos elegido.", "Direccion/contador", "Pendiente"],
            ["Excel/Gymforce entregados.", "Operacion", "Pendiente"],
            ["Reglas mora/descuentos/topes aprobadas.", "Direccion", "Pendiente"],
            ["Plan de capacitacion por rol.", "Operacion/desarrollo", "Pendiente"],
            ["Piloto por sede definido.", "Direccion", "Pendiente"],
            ["Plan de contingencia documentado.", "Operacion", "Pendiente"],
        ], [2.5, 2.0, 2.3]),
        P("Criterios de cierre Sprint 2", s["Heading1"]),
        bullets([
            "La aplicacion corre en staging con datos migrados y usuarios reales.",
            "Existe al menos un piloto operativo por sede o una sede prioritaria.",
            "El contador puede generar reporte semanal sin reconstruirlo manualmente en Excel.",
            "Pagos reales estan integrados o existe decision formal de mantener simulacion/manual temporal.",
            "Backups y restauracion fueron probados.",
            "Permisos fueron revisados para datos de menores y pagos.",
            "Se documentaron riesgos remanentes para Sprint 3.",
        ], s["BodyText"]),
    ]


def main():
    build_pdf(DOCS / "Futsi_Documentacion_Codigo.pdf", code_documentation_story())
    build_pdf(DOCS / "Futsi_Documentacion_Negocio.pdf", business_documentation_story())
    build_pdf(DOCS / "Futsi_Gobernanza_Sprint2_Actualizacion.pdf", governance_update_story())
    governance_main = DOCS / "Futsi_Gobernanza_Roadmap.pdf"
    governance_appendix = DOCS / "Futsi_Gobernanza_Sprint2_Actualizacion.pdf"
    if governance_main.exists() and governance_appendix.exists():
        writer = PdfWriter()
        for source in [governance_main, governance_appendix]:
            reader = PdfReader(str(source))
            for page in reader.pages:
                writer.add_page(page)
        with governance_main.open("wb") as handle:
            writer.write(handle)
    build_docx(
        DOCS / "Futsi_Documentacion_Codigo.docx",
        "Documentacion tecnica del codigo",
        [
            ("Resumen", ["Backend Django/DRF en back/ y frontend React/Vite/Tailwind en front/. API REST con Token Authentication y portales por rol."]),
            ("Estructura", ["- back/core/models.py: dominio.", "- back/core/views.py: API y permisos.", "- back/core/serializers.py: saldos y pagos.", "- front/src/main.tsx: portales y UI.", "- tools/: generadores de docs."]),
            ("Comandos", ["- Backend: manage.py migrate, seed_demo, runserver, test.", "- Frontend: npm install, npm run dev, npm run build."]),
            ("Sprint 2 tecnico", ["- PostgreSQL.", "- S3.", "- Pagos reales.", "- Auditoria automatica.", "- Importador Excel.", "- Refactor frontend."]),
        ],
    )
    build_docx(
        DOCS / "Futsi_Documentacion_Negocio.docx",
        "Documentacion de negocio",
        [
            ("Resumen", ["El sistema controla asistencia, cobranza, descuentos, gastos y reportes por sede para reducir fugas y falta de trazabilidad."]),
            ("Procesos", ["- Academia infantil.", "- Torneos adultos.", "- Cobranza.", "- Gastos.", "- Cierres.", "- Reportes."]),
            ("Reglas pendientes", ["- Mora.", "- Topes para jugar.", "- Reembolsos.", "- Pausas.", "- Descuentos.", "- Inventario uniformes."]),
            ("Sprint 2", ["- Piloto.", "- Migracion.", "- Pagos reales.", "- Auditoria.", "- Cierre diario.", "- Reportes deportivos."]),
        ],
    )
    print("Sprint 2 documentation generated")


if __name__ == "__main__":
    main()
